#!/usr/bin/env python3
"""生成投融资碳核算模块需求规格说明书 v0.61
v0.61 在 v0.6 的数据流转与算法基础上，参照“新建核算任务”修改样例，
将功能需求章节按原型实际页面、按钮、校验、跳转和业务规则细化为
可开发、可测试的操作级说明。"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_COLOR_INDEX

ROOT     = Path("/Users/fangdanyang/Desktop/HUAXIA BANK")
TEMPLATE = ROOT / "需规模版.docx"
OUT      = ROOT / "docs" / "华夏银行绿金系统-投融资碳核算模块_需求规格说明书-v0.61.docx"


# ─── helpers ──────────────────────────────────────────────────────────────────

def _run(p, text, bold=False, size=None, color=None, highlight=None):
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:  r.font.size = Pt(size)
    if bold:  r.bold = True
    if color: r.font.color.rgb = color
    if highlight: r.font.highlight_color = highlight
    return r

def hp(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p

def np(doc, text="", style="Normal"):
    p = doc.add_paragraph(style=style)
    if text: _run(p, text, size=10.5)
    return p

def fp(doc, text):
    p = doc.add_paragraph(style="Normal Indent")
    if text: _run(p, text, size=10.5)
    return p

def lp(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    if text: _run(p, text, size=10.5)
    return p

def pending(doc, text):
    p = doc.add_paragraph(style="Normal")
    _run(p, "【待确认】" + text, highlight=WD_COLOR_INDEX.YELLOW)
    return p

def note(doc, text):
    p = doc.add_paragraph(style="Normal")
    _run(p, "【说明】" + text, color=RGBColor(0x26, 0x74, 0xc1))
    return p

def set_cell_bg(cell, fill="D9EAF7"):
    tcp = cell._tc.get_or_add_tcPr()
    shd = tcp.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcp.append(shd)
    shd.set(qn("w:fill"), fill)

def add_field_table(doc, rows, title_cols=None):
    if title_cols is None:
        title_cols = ["字段名", "类型", "是否必输", "长度", "默认值", "输入限制（或数据字典）", "说明"]
    t = doc.add_table(rows=1, cols=len(title_cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(title_cols):
        hdr[i].text = h
        set_cell_bg(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.name = "宋体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(9)
    return t

def add_plain_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.name = "宋体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(10)
    return t

def add_code_block(doc, lines):
    """伪代码/流程文本块"""
    for line in lines:
        p = doc.add_paragraph(style="Normal Indent")
        r = p.add_run(line)
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9)

def feat(doc, no, name, desc, input_rows, flow_lines, check_lines,
         output_text, ui_text, rule_lines, data_text, output_rows=None):
    heading_name = "新建核算任务-范畴确定" if no == "CR21093-A_ICC_GN0102" else name
    hp(doc, heading_name, 4)
    fp(doc, f"功能编号：{no}")
    fp(doc, f"功能名称：{name}")
    fp(doc, f"功能描述：{desc}")
    fp(doc, "输入项：")
    if input_rows: add_field_table(doc, input_rows)
    if check_lines:
        fp(doc, "校验：")
        for i, c in enumerate(check_lines, 1):
            lp(doc, f"{i}、{c}")
    fp(doc, "处理流程：")
    for i, fl in enumerate(flow_lines, 1):
        lp(doc, f"{i}、{fl}")
    fp(doc, f"输出项：{output_text}")
    if output_rows:
        add_field_table(doc, output_rows, title_cols=["字段名", "类型", "数据种类", "长度/精度"])
    fp(doc, f"界面设计：{ui_text}")
    fp(doc, "凭证报表：不涉及。")
    fp(doc, "会计分录：不涉及。")
    fp(doc, "业务规则：")
    for i, r in enumerate(rule_lines, 1):
        lp(doc, f"{i}、{r}")
    fp(doc, f"数据规则：{data_text}")


# ─── 主函数 ───────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Plain Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CR21093-A_华夏银行绿色金融信息管理系统需求_投融资碳核算模块")
    r.bold = True; r.font.size = Pt(18)
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph(style="Plain Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("需求规格说明书")
    r.bold = True; r.font.size = Pt(22)
    r.font.name = "黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    for line in ["起草人：Cursor", "起草日期：2026.06.29", "文档版本号：v0.61"]:
        p2 = cell.add_paragraph(style="Table_Small")
        _run(p2, line, size=11)
    for p2 in cell.paragraphs:
        if not p2.text.strip():
            p2._element.getparent().remove(p2._element); break

    np(doc, "本文档中所包含的信息属于内部资料，如无华夏银行的书面许可，任何人都无权复制或利用。", style="Normal Indent")

    # ── 修订记录 ──────────────────────────────────────────────────────────────
    hp(doc, "修订记录", 1)
    add_plain_table(doc, ["版本", "日期", "修订描述", "起草/修订人", "复审人"], [
        ["v0.1", "2026.06.10", "依据当前原型生成初稿", "Cursor", ""],
        ["v0.2", "2026.06.10", "参照《需规模版.docx》调整格式，补充输入项/输出项字段表", "Cursor", ""],
        ["v0.3", "2026.06.23", "依据20260617会议纪要全面修订：字段口径更新、业务种类三档终态、截止日期拆分、审核退回逻辑等", "Cursor", ""],
        ["v0.4", "2026.06.24", "对齐最新原型实现：行业范围改为单选组+四级级联、报告法拆分Tab、因子库按组聚合", "Cursor", ""],
        ["v0.5", "2026.06.26", "新增台账管理、基础配置、权限管理；因子库去版本化；数据采集收集状态5档；审核支持调整行业与因子", "Cursor", ""],
        ["v0.6", "2026.06.26",
         "全面补充数据流转逻辑（含端到端流程图、月均余额计算细节、候选→正式字段映射、客户归集合并规则、各核算方法字段全集、归因排放完整公式）；新增基础配置与数据采集联动说明；补充数据采集-客户归集合并改造方案",
         "Cursor", ""],
        ["v0.61", "2026.06.29",
         "参照《需规修改样例-新建核算任务》颗粒度和样式，对4.1功能需求章节按本地原型实际功能细化：补充各功能点输入项、显式编号校验、操作流程、输出项、界面设计、业务规则和数据规则；新建核算任务标题调整为「新建核算任务-范畴确定」",
         "Cursor", ""],
    ])

    hp(doc, "需求对应记录", 1)
    add_plain_table(doc, ["对应需求编号", "版本", "起草/修订人", "说明"], [
        ["CR21093-A", "v0.61", "Cursor",
         "投融资碳核算模块需求规格说明书，v0.61在v0.6数据流转与计算逻辑基础上，按原型实际功能细化功能需求章节"],
    ])
    add_plain_table(doc, ["需求类型", ""], [
        ["需求类型", "□监管部门需求  □产品服务类  ☑管理决策类  □变更维护类  □其他"],
        ["计划情况", "☑本年度需求计划内  □本年度需求计划外"],
    ])

    hp(doc, "目录", 1)
    np(doc, "（目录页码可在 Word 中通过「引用 - 更新目录」生成。）")

    # ──────────────────────────────────────────────────────────────────────────
    # 1 引言
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "1引言", 1)
    hp(doc, "1.1目的", 2)
    np(doc, "本需求规格说明书针对华夏银行绿色金融信息管理系统中的投融资碳核算模块，明确系统应具有的功能、数据处理要求、数据流转规则及计算公式，作为业务人员、开发人员、测试人员和文档编写人员的共同依据。", style="Normal Indent")
    np(doc, "v0.61在v0.6基础上，参照《需规修改样例-新建核算任务》对4.1功能需求章节进行操作级细化：各功能点按原型实际页面补充输入项、校验、处理流程、输出项、界面设计、业务规则和数据规则，便于开发与测试按功能点验收。")

    hp(doc, "1.2背景", 2)
    np(doc, "华夏银行需对投融资业务开展碳排放核算，满足人民银行监管报送和内部绿色金融管理分析需要。投融资碳核算能力以功能模块形式嵌入现有绿金系统，不单独建设独立系统。", style="Normal Indent")
    np(doc, "当前原型已覆盖核算任务发起、接口台账同步、候选清单识别、正式清单确认、核算对象边界确认、数据采集、客户经理在线填报、分级审核、排放计算、结果确认、报告生成、台账管理、排放因子库、企业碳账户及基础配置等完整流程环节。")

    hp(doc, "1.3定义", 2)
    add_plain_table(doc, ["术语", "定义"], [
        ["核算任务",       "围绕某一核算年度或专项范围发起的一次投融资碳核算流程。"],
        ["候选清单",       "系统从信贷台账按规则筛选出的拟纳入核算业务集合，以贷款笔为粒度（1笔=1行）。"],
        ["正式清单",       "管理端确认并锁定后的待核算业务清单，继承候选清单字段，状态推进至confirmed。"],
        ["业务种类",       "三档终态：非项目贷款、项目贷款（以项目方式计算）、项目贷款（以非项目方式计算）。中间状态：项目（计算方法待定）。"],
        ["收集任务",       "针对正式清单派发给客户经理的在线填报任务（supplement），用于补充排放数据、项目信息和佐证附件。"],
        ["归集单元",       "数据采集合并层：同一客户（信用代码）下的同类（非项目/项目）贷款合并为一个收集单元，对应一个supplement。"],
        ["主体排放",       "融资主体或项目在核算周期内的碳排放量（tCO₂e），来源为报告法、能源法、产品法、经济法或兜底法。"],
        ["归因排放",       "按本行投融资占比归因至本行业务的碳排放量（tCO₂e），计算公式见3.3节。"],
        ["月均贷款余额",   "核算年度内该笔贷款存续月份的月末余额均值；内部存储单位为万元，展示单位为元（×10000）。"],
        ["华夏存续月份",   "该笔贷款在核算年度内于华夏银行存续的月份数（1~12），由发放日期与核算年度计算得出。"],
        ["月末余额合计",   "月均贷款余额×华夏存续月份，用于归因分母加总。"],
        ["平均资产总额",   "上年末合并报表资产总额与当年末合并报表资产总额之和除以2，单位元，来源信贷接口。"],
        ["DQR",           "数据质量评级：DQR=Σ(单笔排放×质量得分)/Σ单笔排放，按A/B+/B/B-/C五档。"],
        ["企业碳账户",     "以统一社会信用代码为主键的客户维度排放账户；确认结果后自动归集。"],
        ["信贷数据兜底法", "以月均贷款余额×行业排放因子估算主体排放，对应「其他计算法」，DQR质量等级最低。"],
        ["方法模板",       "按行业+业务类型+核算方法维护的采集字段模板，定义收集填报页展示哪些输入字段及计算公式。"],
        ["参数字段库",     "可复用的字段定义集合，供方法模板引用；包含字段名、数据类型、单位、枚举选项等。"],
    ])

    hp(doc, "1.4参考资料", 2)
    for ref in [
        "《商业银行投融资业务碳核算与报告指南》（人行）",
        "《华夏银行绿金系统-投融资碳核算与企业碳账户-项目需求说明书》v0.1",
        "《华夏银行投融资碳核算及企业碳账户需求评审会议纪要》2026年6月17日、2026年6月23日",
        "《数据采集-客户归集合并改造计划》（内部方案稿，2026.06.23）",
        "当前投融资碳核算原型（assets/js/）及自测脚本（test-flow.js、test-full.js）",
    ]:
        np(doc, ref)

    # ──────────────────────────────────────────────────────────────────────────
    # 2 任务概述
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "2任务概述", 1)
    hp(doc, "2.1目标", 2)
    np(doc, "建设嵌入绿金系统的投融资碳核算模块，支持总行或分行完成从核算任务创建、信贷台账同步、清单确认、数据采集、排放计算、结果确认到报告输出和碳账户归集的完整闭环。")
    for item in [
        "支持按核算年度、行业范围和组织范围发起核算任务；行业范围支持参数化配置。",
        "支持从接口管理同步信贷台账，生成候选清单和正式清单；行业识别以贷款投向所属行业为主要筛选依据。",
        "支持按信贷品种判断项目类/非项目类，结合项目信息形成三档业务种类终态。",
        "数据采集环节支持客户归集合并：同一客户（统一社会信用代码）下的同类贷款合并为一个收集单元下发；填报完成后主体排放拆回逐笔归因。",
        "支持系统自动获取格澜等外部接口数据（报告法来源）和行内系统数据（营收、余额等）形成系统底稿。",
        "支持向客户经理发放数据采集任务；收集填报页字段根据方法模板动态展示，不同行业+方法组合呈现不同字段。",
        "支持分行初审时明确最终采用的核算方法，支持退回至客户经理；审核时可调整归属行业及适用因子。",
        "支持排放计算（含DQR）、结果确认、报告导出和企业碳账户归集。",
        "支持台账管理、基础配置（参数字段库/方法模板/行业配置）和权限管理。",
    ]:
        lp(doc, item)

    hp(doc, "2.2约束", 2)
    for item in [
        "项目类信贷品种完整映射关系尚未由业务侧提供；当前按已知品种判断。",
        "数据采集客户归集合并规则（集团主办行、项目主办分行等）需业务侧确认字段来源，详见4.3节待确认清单。",
        "正式生产环境的接口报文、审批状态码、权限模型、报告模板以华夏银行最终确认材料为准。",
        "审批流、待办中心、组织权限和用户体系不在投融资碳核算模块内重复建设，沿用绿金系统既有能力。",
    ]:
        pending(doc, item)

    # ──────────────────────────────────────────────────────────────────────────
    # 3 端到端数据流转总览（v0.6 新增）
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "3端到端数据流转与核心算法", 1)

    hp(doc, "3.1完整数据流转流程图", 2)
    np(doc, "下图描述投融资碳核算模块从信贷接口到企业碳账户的完整数据流转路径：")
    add_code_block(doc, [
        "┌─────────────────────────────────────────────────────────────────────────────┐",
        "│                    投融资碳核算 — 端到端数据流转                               │",
        "└─────────────────────────────────────────────────────────────────────────────┘",
        "",
        " [信贷核心系统]                                                                 ",
        "  逐月按批次推送 ──→ [接口管理] 按年度汇总成功批次                               ",
        "                                  │                                            ",
        "                         ▼ 点击「从接口同步台账」                                ",
        " [候选清单] ←── 按默认规则筛选（余额≥500万、排除小微/个人/境外/非高碳）           ",
        "    │  每行 = 1笔贷款                                                           ",
        "    │  字段：客户名称、信用代码、贷款账号、信贷品种、一级分行、经办行、            ",
        "    │         月均贷款余额（万元）、月末余额合计（万元）、华夏存续月份、           ",
        "    │         年报营业收入（万元）、合并报表资产总额（万元）、上年末资产（万元）、 ",
        "    │         贷款投向所属行业代码/大类、企业所属行业代码、主办客户经理            ",
        "    │                                                                           ",
        "    ▼ 勾选纳入 → 点击「生成正式清单」                                            ",
        " [正式清单] ← 继承候选清单全部字段，status=draft → confirmed（锁定后）          ",
        "    │  同时：建档企业碳账户空户（status=active）                                 ",
        "    │                                                                           ",
        "    ▼ 点击「调取接口数据」                                                       ",
        "    ├─→ [格澜接口] 按信用代码查询报告法主体排放 → 写入 gelanEntityEmission       ",
        "    │              预填至客户经理填报页「其他来源报告法」Tab（可编辑）            ",
        "    ├─→ [经济法直算] 非项目类 → 主体排放 = 年报营业收入(万元) × 行业因子          ",
        "    │                           归因排放 = 主体排放 × 月均余额 / 平均资产总额    ",
        "    │                                                                           ",
        "    ▼ 「发放收集任务」                                                           ",
        "    ├─→ [客户归集合并] 同一信用代码下的非项目笔合并→1个非项目收集单元             ",
        "    │                  同一信用代码下的项目笔合并→1个项目收集单元                 ",
        "    │                  非项目跨多分行→下发集团主办行；单分行→下发该分行           ",
        "    │                  项目类→下发项目主办分行                                   ",
        "    │                                                                           ",
        "    ▼ [收集任务(supplement)] 1个归集单元=1个supplement                          ",
        "    ├─→ [客户经理填报] 方法模板动态字段                                          ",
        "    │     ├ 报告法-政府/权威来源：scope1/2排放、数据来源、核查情况、附件          ",
        "    │     ├ 报告法-其他来源（格澜预填可编辑）：排放总量、scope1/2、来源           ",
        "    │     ├ 能源法：固定燃料、其他燃料（动态行）、净购入电量、过程排放             ",
        "    │     ├ 产品法：主要产品产量（按模板字段）                                    ",
        "    │     ├ 经济活动法-营收法：只读，系统预填（不可编辑）                         ",
        "    │     └ 项目信息（项目类）：是否可提供→确定业务种类终态                       ",
        "    │                                                                           ",
        "    ▼ 分行审核（初审/终审）                                                      ",
        "    ├─→ 审核通过：选定最终采用方法 → 可调整归属行业与排放因子                     ",
        "    │                                                                           ",
        "    ▼ [排放计算] 主体排放已确定 → 逐笔归因拆分                                   ",
        "    │  非项目：E_归因 = E_主体 × (该笔月均余额) / (归集单元合计平均资产总额)     ",
        "    │  项目（以项目方式）：E_归因 = E_项目 × (该笔月均余额) / 项目总投资         ",
        "    │  兜底法：E_主体 = 月均余额(万) × 行业因子；E_归因 = E_主体（全额归因）     ",
        "    │  DQR 计算：各笔排放加权质量得分→年度综合评级                               ",
        "    │                                                                           ",
        "    ▼ 【确认结果】                                                               ",
        "    └─→ [企业碳账户] 写入排放记录（逐笔）→ 归集至对应账户 annualProfiles         ",
        "                      账户主键 = 统一社会信用代码（同一主体一个账户）             ",
    ])

    hp(doc, "3.2月均贷款余额计算细节", 2)
    np(doc, "候选清单同步时，系统从信贷接口逐笔取月末余额原始数据，按以下规则计算候选清单和正式清单中的余额字段：")
    hp(doc, "3.2.1华夏存续月份", 3)
    np(doc, "华夏存续月份（huaxiaTenureMonths）= 该笔贷款在核算年度内于华夏银行存续的自然月数，取值范围 1~12。", style="Normal Indent")
    add_plain_table(doc, ["场景", "规则", "示例"], [
        ["发放日期早于核算年度", "存续月份 = 12（全年存续）", "2023年发放，核算年度=2024 → 12"],
        ["发放日期在核算年度内", "存续月份 = 12 - 发放月份 + 1", "2024-03月发放 → 12-3+1=10"],
        ["发放日期晚于核算年度", "存续月份 = 1（最小值保护）", "2025年发放，核算年度=2024 → 1"],
        ["发放日期为空", "默认 = 12", "接口未提供发放日期 → 12"],
    ])
    add_code_block(doc, [
        "// 伪代码",
        "function computeHuaxiaTenureMonths(disbursementDate, accountingYear) {",
        "  if (disbursementDate 为空) return 12;",
        "  const [dy, dm] = parseYearMonth(disbursementDate);",
        "  if (dy < accountingYear)  return 12;",
        "  if (dy > accountingYear)  return 1;          // 当年发放但已超年末",
        "  return max(1, 12 - dm + 1);                  // 当年发放",
        "}",
    ])

    hp(doc, "3.2.2月均贷款余额（月均余额）", 3)
    np(doc, "月均贷款余额（avgMonthlyBalance）= 月末余额合计 / 华夏存续月份", style="Normal Indent")
    np(doc, "月末余额合计（monthEndBalanceSum）= 各存续月份月末余额之和。接口未提供逐月月末数据时，以接口提供的avgMonthlyBalance作为初始值：", style="Normal Indent")
    add_code_block(doc, [
        "// 接口提供逐月数据时：",
        "monthEndBalanceSum = sum(月末余额Jan..月末余额Dec [有余额的月份]);",
        "avgMonthlyBalance  = monthEndBalanceSum / huaxiaTenureMonths;",
        "",
        "// 接口仅提供均值时（当前演示场景）：",
        "monthEndBalanceSum = avgMonthlyBalance × huaxiaTenureMonths;",
        "// avgMonthlyBalance 直接用接口值，不再重算",
        "",
        "// 内部存储单位：万元（小数）",
        "// 展示单位：元 → formatLedgerAmountYuan(wan) = wan × 10000 格式化",
    ])
    note(doc, "候选清单中展示「月末余额（元）」= monthEndBalanceSum × 10000；正式清单中展示「月均贷款余额（元）」= avgMonthlyBalance × 10000。两者使用不同字段，切勿混淆。")

    hp(doc, "3.2.3平均资产总额", 3)
    np(doc, "平均资产总额（avgTotalAssets）=（上年末合并报表资产总额 + 当年末合并报表资产总额）/ 2，单位元。", style="Normal Indent")
    add_code_block(doc, [
        "avgTotalAssets = (prevYearTotalAssets + currentYearTotalAssets) / 2;",
        "// prevYearTotalAssets：来源信贷接口，字段名 prevYearTotalAssets（万元）",
        "// currentYearTotalAssets：来源信贷接口，字段名 totalAssets（万元）",
        "// 候选清单展示「合并报表资产总额（元）」= currentYearTotalAssets × 10000",
        "// 正式清单展示「平均资产总额（元）」   = avgTotalAssets × 10000",
    ])

    hp(doc, "3.3候选清单→正式清单字段映射", 2)
    np(doc, "点击「生成正式清单」时，系统按1:1从候选清单创建正式清单记录，字段继承规则如下：")
    add_plain_table(doc, ["字段", "候选清单展示名", "正式清单展示名", "存储字段", "单位/换算", "备注"], [
        ["客户名称",           "客户名称",               "客户名称",               "customerName",          "文本",      ""],
        ["统一社会信用代码",   "统一社会信用代码",         "统一社会信用代码",         "creditCode",            "文本",      "企业碳账户主键"],
        ["贷款账号",           "贷款账号",               "贷款账号",               "loanAccount",           "文本",      ""],
        ["信贷品种",           "信贷品种",               "信贷品种",               "loanType",              "文本",      ""],
        ["业务种类",           "业务种类",               "业务种类",               "accountingType",        "枚举",      "由信贷品种+项目信息推导"],
        ["一级分行",           "一级分行",               "一级分行",               "tier1Branch",           "文本",      ""],
        ["经办行",             "经办行",                 "经办行",                 "handlingBranch",        "文本",      ""],
        ["主办客户经理",       "主办客户经理",             "主办客户经理",             "manager",               "文本",      ""],
        ["华夏存续月份",       "（内部字段）",             "（内部字段）",             "huaxiaTenureMonths",    "整数/月",   "候选同步时计算"],
        ["月末余额合计",       "月末余额（元）",           "（内部参考）",             "monthEndBalanceSum",    "万元→展示元", "候选清单显示此字段"],
        ["月均贷款余额",       "（内部字段）",             "月均贷款余额（元）",       "avgMonthlyBalance",     "万元→展示元", "正式清单显示此字段"],
        ["年报营业收入",       "年报营业收入（元）",       "年报营业收入（元）",       "operatingRevenue",      "万元→展示元", ""],
        ["上年末合并报表资产", "（内部字段）",             "（内部字段）",             "prevYearTotalAssets",   "万元",      "用于计算avgTotalAssets"],
        ["当年末合并报表资产", "合并报表资产总额（元）",   "（内部参考）",             "totalAssets",           "万元→展示元", "候选清单显示此字段"],
        ["平均资产总额",       "（内部字段）",             "平均资产总额（元）",       "avgTotalAssets",        "万元→展示元", "正式清单显示此字段"],
        ["贷款投向行业代码",   "贷款投向所属行业",         "贷款投向所属行业",         "gbIndustryCode",        "文本/GB四级", "筛选主依据"],
        ["企业所属行业代码",   "企业所属行业",             "企业所属行业",             "industryMajor",         "文本",      "辅助展示字段"],
        ["项目明细",           "（展开查看）",             "（展开查看）",             "projectDetails",        "JSON数组",  "项目类业务"],
    ])

    hp(doc, "3.4数据采集客户归集合并规则", 2)
    np(doc, "正式清单锁定后，「发放收集任务」前系统自动执行客户归集合并（buildCollectGroups），将同一客户的多笔贷款合并为归集单元（collectGroup），以便按客户维度统一下发收集任务、填报和审核。")
    note(doc, "台账逐笔数据不动；合并只发生在「数据采集下发」层。归因排放仍逐笔计算，以保证监管报送精度。")

    hp(doc, "3.4.1归集维度", 3)
    add_plain_table(doc, ["归集维度", "规则"], [
        ["客户主键",     "统一社会信用代码（creditCode），同一主体不同贷款账号合并"],
        ["归集桶",       "非项目（non_project）和项目（project）分开，同一客户最多2个收集单元"],
        ["桶判定",       "非项目贷款→非项目桶；项目类（含待定）→项目桶"],
    ])

    hp(doc, "3.4.2下发分行规则", 3)
    add_plain_table(doc, ["场景", "下发对象", "说明"], [
        ["非项目，仅1个一级分行",    "该分行",       "单分行直接下发"],
        ["非项目，跨多个一级分行",   "集团主办行",    "需接口/CRM提供groupLeadBranch字段；待确认缺失时默认规则"],
        ["项目类",                   "项目贷款主办分行", "需字段projectLeadBranch；多项目跨分行时优先级待确认"],
    ])
    pending(doc, "集团主办行、项目贷款主办分行字段的来源（信贷接口/CRM/人工指定）及缺失时默认值，待业务与技术确认。")

    hp(doc, "3.4.3金额字段归集方式", 3)
    add_plain_table(doc, ["字段", "归集方式", "说明"], [
        ["月均贷款余额",       "求和（Σ各笔）",                 "归集单元合计余额，用于后续逐笔归因分母"],
        ["年报营业收入",       "客户级去重（取首笔或最大值）",   "合并报表口径，不重复累加"],
        ["平均资产总额",       "客户级去重",                     "合并报表口径，不重复累加"],
        ["项目类字段",         "保留projectDetails列表",         "多项目合并至同一收集单元"],
    ])

    hp(doc, "3.4.4归集核算行业规则（非项目）", 3)
    add_code_block(doc, [
        "// 非项目归集单元核算行业确定规则：",
        "const investCodes = unique(nonProject.map(r => r.gbIndustryCode));",
        "if (investCodes.length > 1) {",
        "  // 多个投向行业 → 取客户所属行业（企业所属行业，industryMajor）",
        "  accountingIndustryCode   = pickCustomerIndustry(creditCode, nonProject);",
        "  accountingIndustrySource = 'customer';",
        "} else {",
        "  // 单一投向 → 用贷款投向行业",
        "  accountingIndustryCode   = investCodes[0];",
        "  accountingIndustrySource = 'invest';",
        "}",
        "// 项目类 → 仍按项目明细各自行业填报，不做行业合并",
    ])
    pending(doc, "非项目单一投向时，核算行业使用投向还是客户所属行业，请业务侧确认统一口径。")

    hp(doc, "3.4.5归集单元数据结构", 3)
    add_code_block(doc, [
        "collectGroup = {",
        "  id: 'G{taskId}_{seq}',",
        "  taskId,",
        "  creditCode,",
        "  customerName,",
        "  bucket: 'non_project' | 'project',",
        "  memberFormalIds: ['F001', 'F002', ...],  // 关联逐笔正式清单",
        "  memberCount: 3,",
        "  dispatchBranch: '上海分行',",
        "  dispatchRule: 'single_branch' | 'group_lead' | 'project_lead',",
        "  accountingIndustryCode: 'C3011',",
        "  accountingIndustrySource: 'customer' | 'invest' | 'project',",
        "  aggregatedBalance: 12500,              // 合计月均余额（万元）",
        "  revenue: 68000,                        // 客户级营业收入（万元，去重）",
        "  totalAssets: 320000,                   // 客户级平均资产总额（万元，去重）",
        "  groupLeadBranch: '北京分行',",
        "  projectLeadBranch: '深圳分行',",
        "  status: 'pending' | 'dispatched' | 'completed',",
        "  supplementId: 'S...'                   // 1:1关联supplement",
        "}",
    ])

    hp(doc, "3.5各核算方法采集字段全集", 2)
    np(doc, "客户经理在线收集填报页（supplement）中，根据行业+业务类型+方法模板动态展示字段。以下按方法类型列出所有可能采集字段（具体行业模板可能为子集）：")

    hp(doc, "3.5.1报告法—政府/权威来源（reportAuthority）", 3)
    add_field_table(doc, [
        ("报告法数据来源", "下拉框", "是", "50", "EPD认证报告", "ESG报告/环境信息披露/政府核查/CCER/第三方核查", ""),
        ("该数据是否经政府/第三方核查", "下拉框", "是", "10", "是", "是/否", "是→须上传附件"),
        ("温室气体排放总量（tCO₂e）", "数字输入", "是", "20,4", "", ">0", "scope1+scope2合计"),
        ("范围一排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", "直接排放"),
        ("范围二排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", "间接排放"),
        ("核查报告期间起", "日期", "否", "10", "", "YYYY-MM-DD", ""),
        ("核查报告期间止", "日期", "否", "10", "", "YYYY-MM-DD", ""),
        ("报告附件", "文件上传", "条件必填", "—", "", "pdf/doc/xls/png/jpg，最多3个，单文件≤20MB", "核查=是时必传"),
    ])

    hp(doc, "3.5.2报告法—其他来源（reportOther，含格澜预填）", 3)
    add_field_table(doc, [
        ("报告法数据来源", "下拉框", "是", "50", "格澜数据-各地区企业环境信息披露平台", "ESG报告/格澜数据-.../行业协会/其他", "格澜数据预填此字段"),
        ("该数据是否经政府/第三方核查", "下拉框", "是", "10", "否", "是/否", ""),
        ("温室气体排放总量（tCO₂e）", "数字输入", "是", "20,4", "", ">0", "格澜预填，可编辑"),
        ("范围一排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", ""),
        ("范围二排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", ""),
        ("报告附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    note(doc, "格澜接口成功返回数据后，系统自动将数据写入「其他来源报告法」Tab，并标注来源为「格澜数据-各地区企业环境信息披露平台」。客户经理可在填报页查看并修改数值或来源说明，但不可删除该预填记录。")

    hp(doc, "3.5.3物理活动法—能源法（energy，字段由方法模板决定）", 3)
    np(doc, "能源法字段分为「固定燃料」「其他燃料（动态行）」「净购入电量/热力」「过程排放（动态行）」四个部分，具体字段由对应行业的方法模板配置决定。以下列出系统内已实现的字段集合：")
    add_field_table(doc, [
        ("固定燃料-原煤消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", "方法模板fuelFixed"),
        ("固定燃料-焦炭消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-燃料油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-天然气消耗量（万立方米）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-汽油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-柴油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("其他燃料（动态行）-燃料品种", "下拉框", "否", "50", "", "原油/燃料油/汽油/柴油/液化石油气/液化天然气/焦炉煤气/高炉煤气/转炉煤气/炼厂干气/其他煤气/无", "可增删行，最多20行"),
        ("其他燃料（动态行）-消耗量（吨或万立方米）", "数字输入", "否", "20,4", "", ">0", "与品种行对应"),
        ("所属电网", "下拉框", "是", "50", "全国平均", "全国平均/华北/华东/华中/南方/东北/西北", "由模板gridOptions决定"),
        ("净购入电量（MWh）", "数字输入", "是", "20,4", "", "≥0", ""),
        ("净购入热力（GJ）", "数字输入", "否", "20,4", "", "≥0", "仅部分模板展示"),
        ("过程排放-脱硫试剂类型", "下拉框", "否", "50", "", "石灰石/白云石/方解石/其他", "动态行，按模板processBlocks"),
        ("过程排放-脱硫试剂消耗量（吨）", "数字输入", "否", "20,4", "", ">0", ""),
        ("过程排放-碳酸盐类型", "下拉框", "否", "50", "", "石灰石/白云石/其他", "动态行，按模板"),
        ("过程排放-碳酸盐消耗量（吨）", "数字输入", "否", "20,4", "", ">0", ""),
        ("能源法附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    np(doc, "能源法主体排放计算公式（演示原型暂不自动试算，以业务确认版为准）：")
    add_code_block(doc, [
        "E_主体 = Σ(燃料消耗量 × 对应排放因子)         // 燃料燃烧",
        "       + Σ(脱硫/碳酸盐消耗量 × 工艺因子)       // 过程排放",
        "       + 净购入电量(MWh) × 区域电网因子         // 外购电力",
        "       + 净购入热力(GJ) × 热力因子;             // 外购热力（部分行业）",
    ])

    hp(doc, "3.5.4物理活动法—产品法（product，字段由方法模板决定）", 3)
    np(doc, "产品法字段完全由对应行业+方法模板的 product.fields 配置决定，不同行业字段完全不同。以下列出通用字段结构：")
    add_field_table(doc, [
        ("主要产品名称/类型", "文本/模板内置", "按模板", "100", "", "", "模板中固定展示"),
        ("主要产品产量（万吨）", "数字输入", "是", "20,4", "", ">0", "按模板分组展示"),
        ("细分产品产量（万吨）", "数字输入", "按模板", "20,4", "", ">0", "部分行业有细分"),
        ("产品法附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    add_code_block(doc, [
        "E_主体 = Σ(产品产量 × 产品碳排放因子);   // 各产品类型分别计算后求和",
        "// 产品碳排放因子来源：排放因子库 caliberTag=pbo（人行口径）或 caliberTag=bank（行自定义）",
    ])

    hp(doc, "3.5.5经济活动法—营收法（economy，系统预填只读）", 3)
    np(doc, "经济活动法-营收法字段由系统从正式清单数据预填，客户经理填报页只读，不可编辑。")
    add_field_table(doc, [
        ("年报营业收入（万元）", "只读展示", "—", "20,2", "", "来自正式清单operatingRevenue", "系统自动填入"),
        ("行业排放因子（tCO₂e/万元营收）", "只读展示", "—", "20,6", "", "来自因子库经济活动法行业因子", "系统自动匹配"),
        ("系统计算主体排放（tCO₂e）", "只读展示", "—", "20,4", "", "= 营收 × 因子", "系统计算"),
    ])
    add_code_block(doc, [
        "E_主体 = 年报营业收入(万元) × 行业排放因子(tCO₂e/万元营收);",
        "// 行业因子匹配优先级：GB国标四级代码 > 行业大类",
        "// 若营收为0或缺失 → 该方法无法计算，提示「请提供营业收入」",
    ])

    hp(doc, "3.5.6信贷数据兜底法（credit_fallback，不在收集填报页展示）", 3)
    np(doc, "信贷数据兜底法由管理端直接执行，不在客户经理收集填报页展示。触发条件：仍有记录无主体排放时，点击「信贷数据兜底法」按钮。")
    add_code_block(doc, [
        "E_主体 = 月均贷款余额(万元) × 行业排放因子(tCO₂e/万元余额);",
        "E_归因 = E_主体;   // 全额归因，不按资产占比折算",
        "// 核算方法标记为「其他计算法」；DQR质量等级=5（最低）",
        "// 行业因子取 gbIndustryCode 对应经济活动法因子",
    ])

    hp(doc, "3.6归因排放完整计算公式", 2)
    np(doc, "排放计算环节，系统按业务种类和核算路径分档计算归因排放：")
    add_plain_table(doc, ["业务种类", "主体排放来源", "归因排放公式", "归因占比分母", "说明"], [
        ["非项目贷款",
         "报告法/能源法/产品法/经济法/兜底法",
         "E_归因 = E_主体 × (该笔月均余额/12) / (平均资产总额×12) × 12",
         "融资主体平均资产总额",
         "简化写法：E_归因 = E_主体 × 月均余额 / 平均资产总额"],
        ["项目贷款（以项目方式计算）",
         "客户经理填报项目排放数据",
         "E_归因 = E_项目 × 该笔月均余额 / 项目总投资",
         "项目总投资（projectTotalInvestmentWan×10000）",
         "项目维度归因"],
        ["项目贷款（以非项目方式计算）",
         "同非项目贷款",
         "E_归因 = E_主体 × 月均余额 / 平均资产总额",
         "融资主体平均资产总额",
         "与非项目同公式"],
        ["信贷数据兜底法",
         "月均余额×行业因子",
         "E_归因 = E_主体（全额归因）",
         "不折算",
         "DQR最低"],
    ])
    add_code_block(doc, [
        "// 伪代码",
        "for each formal f in confirmedFormals(taskId) {",
        "  const calc = getCalculation(f.id);",
        "  if (calc.source === 'credit_fallback') {",
        "    f.attributedEmission = calc.entityEmission;   // 全额归因",
        "  } else if (f.accountingType === 'project_as_project') {",
        "    const projectInvest = f.projectTotalInvestmentWan * 10000;",
        "    f.attributedEmission = calc.entityEmission * f.avgMonthlyBalance / projectInvest;",
        "  } else {",
        "    // non_project 或 project_as_non_project",
        "    const avgAssets = f.avgTotalAssets * 10000;   // 换算为元",
        "    const balance   = f.avgMonthlyBalance * 10000; // 换算为元",
        "    f.attributedEmission = calc.entityEmission * balance / avgAssets;",
        "  }",
        "}",
    ])

    hp(doc, "3.7方法模板与参数字段库联动机制", 2)
    np(doc, "基础配置模块（参数字段库+方法模板）与数据采集收集填报的联动流程如下：")
    add_code_block(doc, [
        "配置层（基础配置/方法模板）：",
        "  总行维护「参数字段库」→ 定义可复用字段（名称/类型/单位/枚举等）",
        "  总行维护「方法模板」→ 按行业+业务类型+核算方法组合，",
        "                         引用参数字段库中的字段，配置计算公式和因子绑定",
        "  点击「发布」→ 模板status=published，对应行业+方法组合生效",
        "",
        "展示层（数据采集/收集填报）：",
        "  系统根据 formal.industryMajor + formal.accountingType",
        "  查找 status=published 的方法模板",
        "  若找到 → 按模板字段渲染收集填报页各方法Tab",
        "  若未找到 → 按默认字段集（defaultMode）渲染",
        "",
        "字段渲染规则（以能源法为例）：",
        "  template.methods.energy.fuelFixed  → 固定燃料输入行",
        "  template.methods.energy.otherFuelOptions → 其他燃料下拉选项",
        "  template.methods.energy.gridOptions → 所属电网下拉选项",
        "  template.methods.energy.hasPurchasedHeat → 是否展示净购入热力字段",
        "  template.methods.energy.processBlocks → 过程排放动态行配置",
        "  template.methods.product.fields → 产品法产品字段列表",
    ])
    note(doc, "当前原型已实现玻璃行业、能源行业等方法模板示例，其余行业模板由总行管理员在「方法模板配置」页面维护后发布生效，无需修改系统代码。")

    hp(doc, "3.8主体排放→逐笔归因拆分流程", 2)
    np(doc, "客户归集合并后，supplement填报的主体排放为归集单元（collectGroup）级别，需在排放计算时拆回逐笔归因：")
    add_code_block(doc, [
        "// 归集单元级主体排放已由审核确认",
        "const groupEntityEmission = supplement.fieldData.entityEmission;",
        "",
        "for each formalId in collectGroup.memberFormalIds {",
        "  const f = getFormal(formalId);",
        "  // 非项目：各笔共享同一主体排放，按各笔余额占比归因",
        "  const totalBalance = sum(memberFormals.map(m => m.avgMonthlyBalance));",
        "  f.entityEmission    = groupEntityEmission;",
        "  f.attributedEmission = groupEntityEmission",
        "                        * f.avgMonthlyBalance / f.avgTotalAssets;",
        "  // 注：avgTotalAssets为客户级（去重），各笔相同",
        "}",
    ])

    hp(doc, "3.9DQR数据质量评价", 2)
    add_plain_table(doc, ["核算方法", "质量得分（qualityScore）", "说明"], [
        ["报告法-政府/权威来源（经核查）", "1", "最高质量"],
        ["报告法-政府/权威来源（未核查）", "2", ""],
        ["报告法-其他来源（含格澜）",      "2", ""],
        ["物理活动法-能源法",              "2", ""],
        ["物理活动法-产品法",              "3", ""],
        ["经济活动法-营收法（直算）",      "4", ""],
        ["信贷数据兜底法（其他计算法）",   "5", "最低质量"],
    ])
    add_code_block(doc, [
        "// 年度综合DQR",
        "DQR = Σ(E_归因_i × qualityScore_i) / Σ(E_归因_i);",
        "// DQR等级映射：",
        "// ≤1.5 → A     ≤2.0 → B+    ≤3.0 → B",
        "// ≤4.0 → B-    >4.0 → C",
    ])

    # ──────────────────────────────────────────────────────────────────────────
    # 4 功能需求（与v0.5一致，引用3章中的数据规则）
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "4需求规定", 1)
    hp(doc, "4.1功能需求", 2)
    np(doc, "以下各功能点中的数据规则、计算公式和字段映射，以第3章为权威说明；本章侧重输入/输出规格和处理流程。")

    # ── 4.1.1 核算任务管理 ─────────────────────────────────────────────────
    hp(doc, "4.1.1核算任务管理", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0101",
        name="查询核算任务列表",
        desc="用于查询当前用户有权访问的投融资碳核算任务，并作为进入后续流程步骤的入口。",
        input_rows=[
            ("任务名称", "文本框", "否", "100", "", "", "模糊匹配"),
            ("核算年度", "下拉框", "否", "4", "", "", ""),
            ("投向行业范围", "下拉框", "否", "50", "", "全部/八大高碳/八大+扩展/自定义", ""),
            ("所属行业范围", "下拉框", "否", "50", "", "同上", ""),
            ("当前进度", "下拉框", "否", "20", "", "范畴确定/清单识别/对象边界/数据采集/排放计算/生成报告", ""),
            ("分页参数", "系统参数", "是", "—", "", "", ""),
        ],
        flow_lines=[
            "用户进入核算任务管理菜单。",
            "系统根据当前角色和组织权限过滤任务列表，分行角色仅显示本分行发起或涉及本分行的任务。",
            "按筛选条件过滤，展示任务名称、核算年度、投向行业范围、所属行业范围、数据采集截止日期、分行审批截止日期、当前进度及操作列。",
            "操作列：【编辑】【查看】【删除】；工具栏：【+新建核算任务】。",
        ],
        check_lines=["客户经理角色不可访问核算任务管理菜单。"],
        output_text="任务列表、分页信息。",
        output_rows=[
            ("任务ID", "文本", "字符", "50"),
            ("任务名称", "文本", "字符", "100"),
            ("核算年度", "文本", "数值", "4"),
            ("投向行业范围", "文本", "字符", "50"),
            ("所属行业范围", "文本", "字符", "50"),
            ("数据采集截止日期", "文本", "日期", "10"),
            ("分行审批截止日期", "文本", "日期", "10"),
            ("当前进度", "文本", "字符", "20"),
        ],
        ui_text="核算任务管理列表页，顶部工具栏含【+新建核算任务】按钮，下方为筛选区域和分页表格。",
        rule_lines=[
            "总行和分行按系统权限查看任务；客户经理不展示该菜单。",
            "六步流程条：范畴确定→清单识别→对象边界→数据采集→排放计算→生成报告，当前进度对应workflowStep枚举值。",
        ],
        data_text="任务字段：id、name、year、subjectIndustryScope、investIndustryScope、deadline、branchDeadline、workflowStep、initiatorOrg、branches（组织范围数组）。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0102",
        name="新建核算任务",
        desc="用于创建新的投融资碳核算任务，设置核算范围参数，并启动清单识别流程。",
        input_rows=[
            ("任务名称", "输入域", "是", "100", "", "", ""),
            ("核算年度", "数字输入+步进器", "是", "4", "当前年", "2020~2030", ""),
            ("数据行业范围", "单选", "是", "—", "投向行业范围", "投向行业范围/所属行业范围", "决定候选清单筛选依据"),
            ("行业范围（投向/所属）", "单选组", "是", "50", "人行投融资碳核算八大高碳行业", "人行八大高碳/人行八大高碳+我行主要行业/自定义", ""),
            ("行业自定义（投向/所属）", "四级级联多选", "条件必填", "—", "", "GB/T4754-2017四级行业", "行业范围=自定义时展开"),
            ("余额口径", "下拉框", "是", "50", "月均余额", "月均余额/日均余额", ""),
            ("组织范围", "多选框", "是", "—", "全行", "全行/各一级分行名称", "支持多选"),
            ("数据采集截止日期", "日期", "是", "10", "", "YYYY-MM-DD", ""),
            ("分行审批截止日期", "日期", "否", "10", "", "YYYY-MM-DD", "非强制截止"),
        ],
        flow_lines=[
            "点击【+新建核算任务】进入新建页，顶部展示六步流程条（当前步=第一步：范畴确定）。",
            "选择「数据行业范围」（投向行业范围/所属行业范围），切换后展示对应行业范围单选组。",
            "行业范围单选组三档：人行投融资碳核算八大高碳行业、人行八大高碳+我行主要行业、自定义；默认人行八大高碳。",
            "「自定义」时展开GB/T4754-2017四级级联多选面板（支持行业全选/清空/按代码搜索）。",
            "组织范围多选：全行或多个一级分行（不可两者都选，选全行时分行自动禁用）。",
            "点击【保存并启动】→ 校验通过 → 创建任务，workflowStep推进至清单识别 → 跳转候选业务清单页。",
        ],
        check_lines=[
            "任务名称、核算年度（2020~2030）、行业范围、余额口径、组织范围、数据采集截止日期为必填。",
            "「自定义」行业时须至少选择一个行业代码；否则提示「请至少选择一项行业」。",
            "组织范围须至少选择全行或一个一级分行。",
        ],
        output_text="新建任务记录；currentTaskId更新；跳转候选清单页。",
        output_rows=[
            ("任务ID", "文本", "字符", "50"),
            ("workflowStep", "枚举", "数值", "1"),
        ],
        ui_text="新建核算任务页，六步流程条+两栏网格表单；行业范围单选组+四级级联面板；底部含【取消】【保存并启动】。",
        rule_lines=[
            "行业范围三档文案：人行投融资碳核算八大高碳行业 / 人行八大高碳+我行主要行业 / 自定义。",
            "行业范围对应的实际行业代码集合由「基础配置/行业配置」模块中的标识（pbo/bank）决定，参见4.1.11节。",
            "总行发起：收集数据审核路径=分行初审→总行终审；分行发起：分行审核通过即完成。",
        ],
        data_text="initiatorOrg（hq/branch）决定审核路径；orgScope=branches数组；industryCodes和investIndustryCodes存储实际行业代码清单。",
    )

    # ── 4.1.2 接口管理与候选清单 ────────────────────────────────────────────
    hp(doc, "4.1.2接口管理与候选清单识别", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0201",
        name="查询接口批次",
        desc="用于查询信贷台账月度同步批次状态，支撑候选清单识别。",
        input_rows=[
            ("数据月份", "日期选择", "否", "7", "", "", ""),
            ("批次状态", "下拉框", "否", "20", "", "全部/成功/失败/获取中", ""),
        ],
        flow_lines=[
            "用户进入接口管理页。",
            "展示统计卡片（总批次数/成功/失败/总台账条数）；列表展示批次号、数据年份、数据月份、推送时间、数据条数、来源系统、状态。",
            "点击【查看】弹窗预览该批次台账字段（字段与候选清单台账字段一致）。",
            "失败批次可点击【重新获取】触发重试。",
        ],
        check_lines=["客户经理不可访问接口管理菜单。"],
        output_text="接口批次列表。",
        output_rows=[("批次号", "文本", "字符", "50"), ("数据月份", "文本", "字符", "7"), ("批次状态", "文本", "字符", "20")],
        ui_text="接口管理列表页；统计卡片行+分页表格。",
        rule_lines=["接口批次与业务核算任务数据隔离；每月1日01:00自动同步上月台账。"],
        data_text="【待确认】生产接口字段、频率、错误码、重试机制、脱敏要求。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0202",
        name="同步候选业务清单",
        desc="从接口管理按核算年度汇总信贷台账，按默认规则筛选形成候选业务清单，并计算月均余额等衍生字段。",
        input_rows=[
            ("任务ID", "系统参数", "是", "50", "", "", ""),
            ("核算年度", "系统参数", "是", "4", "", "", "由当前任务决定"),
        ],
        flow_lines=[
            "用户在候选业务清单页点击【从接口同步台账（X年度）】。",
            "系统汇总该核算年度全部成功批次的台账数据（跨月去重合并同一贷款账号）。",
            "对每笔贷款执行字段归一化：计算huaxiaTenureMonths（见3.2.1节）→ 计算monthEndBalanceSum和avgMonthlyBalance → 计算avgTotalAssets（见3.2.3节）。",
            "按默认规则标记excluded/included（排除规则：月均余额<500万、小微企业、个人个体、境外主体、非高碳行业）。",
            "展示同步时间、接入总笔数、已纳入笔数，初始化候选清单列表。",
        ],
        check_lines=["若无该年度成功批次，提示「接口管理中暂无X年度台账批次，请等待每月1日自动推送」。"],
        output_text="候选业务清单；同步提示。",
        output_rows=[
            ("候选记录ID", "文本", "字符", "50"),
            ("月均贷款余额（元）", "展示值", "数值", "20"),
            ("月末余额（元）", "展示值", "数值", "20"),
            ("平均资产总额（元）", "展示值", "数值", "20"),
        ],
        ui_text="候选业务清单页工具栏含【从接口同步台账】【生成正式清单（n笔）】，下方为筛选面板和分页表格。",
        rule_lines=[
            "候选清单表头：月末余额（元）= monthEndBalanceSum × 10000；合并报表资产总额（元）= totalAssets × 10000。",
            "正式清单表头：月均贷款余额（元）= avgMonthlyBalance × 10000；平均资产总额（元）= avgTotalAssets × 10000。",
            "内部存储单位均为万元（小数），展示时×10000。",
        ],
        data_text="候选记录含：customerName、creditCode、loanAccount、loanType、tier1Branch、handlingBranch、manager、gbIndustryCode（投向）、industryMajor（企业所属）、avgMonthlyBalance、monthEndBalanceSum、huaxiaTenureMonths、operatingRevenue、totalAssets、prevYearTotalAssets、avgTotalAssets、bizType、projectDetails。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0203",
        name="候选清单筛选与纳入管理",
        desc="通过多维度筛选条件精细化管理候选清单，并勾选拟纳入的业务记录。",
        input_rows=[
            ("信贷品种", "多选框", "否", "—", "", "贷款品种枚举", ""),
            ("贷款主体类型", "多选框", "否", "—", "", "有限公司/股份公司/国有企业等", ""),
            ("客户规模", "多选框", "否", "—", "", "大型/中型/小微企业", ""),
            ("投向行业", "多选框", "否", "—", "", "贷款投向所属行业四级代码", ""),
            ("月均贷款余额（元）起", "数字输入", "否", "20", "", "", ""),
            ("月均贷款余额（元）止", "数字输入", "否", "20", "", "", ""),
        ],
        flow_lines=[
            "设置筛选条件后点击【查询】，系统按条件刷新候选清单列表。",
            "逐条或通过全选复选框勾选拟纳入记录（included=true）。",
            "点击【恢复默认筛选条件】重置为按指引核算范畴的默认筛选并自动重新标记included。",
            "点击【清除全部筛选条件】清空所有筛选项（不改变included状态）。",
        ],
        check_lines=[],
        output_text="过滤后的候选清单；已纳入笔数统计。",
        output_rows=[
            ("客户名称", "文本", "字符", "100"),
            ("信贷品种", "文本", "字符", "100"),
            ("业务种类", "文本", "字符", "50"),
            ("月末余额（元）", "数值", "数值", "20"),
            ("合并报表资产总额（元）", "数值", "数值", "20"),
            ("年报营业收入（元）", "数值", "数值", "20"),
            ("主办客户经理", "文本", "字符", "100"),
        ],
        ui_text="候选业务清单页，筛选面板，表格含多选复选框，底部统计卡片。",
        rule_lines=[
            "候选清单筛选以贷款投向所属行业（gbIndustryCode）为主要依据，企业所属行业（industryMajor）保留为辅助展示字段。",
            "投放金额字段不在候选/正式清单列表中展示。",
        ],
        data_text="排除规则枚举：LOW_BALANCE（余额<500万）、SME、INDIVIDUAL、OVERSEAS、NON_HIGH_CARBON；candidateFilterRules.customized标记是否手动调整。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0204",
        name="业务种类识别",
        desc="在候选清单、正式清单和数据采集列表中展示业务种类，并在客户经理填报项目信息后完成终态定档。",
        input_rows=[
            ("信贷品种", "系统字段", "是", "100", "", "", ""),
            ("项目明细", "系统字段", "否", "—", "", "", "projectDetails是否非空"),
            ("是否可提供项目信息", "下拉框", "否", "10", "", "是/否", "客户经理填报时声明"),
        ],
        flow_lines=[
            "系统首先根据信贷品种判断项目类或非项目类：项目贷款/一般性固定资产贷款/出口退税账户托管贷款等→项目类；其余→非项目类。",
            "非项目类 → accountingType=non_project（终态）。",
            "项目类且接口已同步项目明细（projectDetails非空）→ accountingType=project_as_project（终态）。",
            "项目类但无项目明细且客户经理未声明 → accountingType=null（中间状态，展示「项目（计算方法待定）」）。",
            "客户经理填报中选择「否」（不能提供项目信息）→ project_as_non_project（终态）。",
            "客户经理填报中选择「是」并提交项目信息 → project_as_project（终态）。",
            "进入排放计算阶段：剩余「计算方法待定」强制定档为project_as_non_project。",
        ],
        check_lines=["业务种类由系统规则自动推导；普通用户不可手工编辑accountingType字段。"],
        output_text="accountingType（业务种类枚举）；业务种类展示文本。",
        output_rows=None,
        ui_text="「信贷品种」列之后紧接「业务种类」列；数据采集列表同样展示。",
        rule_lines=["【待确认】三档终态名称是否与人行附件4完全一致（含是否带「贷款」二字），请业务侧确认。"],
        data_text="accountingType枚举：non_project、project_as_project、project_as_non_project、null（中间）。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0205",
        name="生成正式清单",
        desc="将勾选纳入的候选业务转为正式核算清单，继承所有台账字段。",
        input_rows=[
            ("候选记录勾选状态", "复选框", "是", "—", "", "", ""),
            ("任务ID", "系统参数", "是", "50", "", "", ""),
        ],
        flow_lines=[
            "勾选拟纳入候选记录，点击【生成正式清单（n笔）】。",
            "系统创建正式清单记录，继承候选记录全部台账字段（含余额、营收、资产、行业、项目明细等，参见3.3节字段映射表）。",
            "正式清单记录status=draft，workflowStep维持在清单识别，跳转至正式清单确认页。",
        ],
        check_lines=[
            "至少存在一笔included=true的候选记录。",
            "同一候选记录不得重复生成正式清单。",
        ],
        output_text="正式清单记录（status=draft）。",
        output_rows=None,
        ui_text="候选清单工具栏【生成正式清单】按钮，括号内显示已勾选笔数。",
        rule_lines=["正式清单继承候选清单所有字段，包含衍生字段（avgMonthlyBalance、avgTotalAssets等）；不重新计算。"],
        data_text="formalList记录：继承candidateId及所有台账字段；accountingType继承；projectDetails继承；status=draft。",
    )

    # ── 4.1.3 正式清单 ──────────────────────────────────────────────────────
    hp(doc, "4.1.3正式清单确认与核算对象边界", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0301",
        name="正式清单查询与锁定",
        desc="展示正式待核算业务清单，支持查看状态与项目明细，并支持锁定操作以推进至数据采集环节。",
        input_rows=[
            ("任务ID", "系统参数", "是", "50", "", "", ""),
            ("勾选的正式清单ID", "复选框", "是", "—", "", "", ""),
        ],
        flow_lines=[
            "进入正式清单确认页，展示业务台账字段（月均贷款余额/平均资产总额/年报营业收入）、业务种类、状态。",
            "项目类业务含子明细时可展开查看项目信息（projectDetails）。",
            "用户勾选或全选记录后点击【确认锁定】。",
            "系统将选中记录status改为confirmed；任务workflowStep推进至数据采集；同步为已锁定记录建档企业碳账户（空户，status=active）。",
            "企业碳账户主键=统一社会信用代码（creditCode），同一信用代码已有账户则不重复建档。",
        ],
        check_lines=["锁定后不可在普通页面直接回退，变更须通过审核退回或管理员操作。"],
        output_text="锁定后的正式清单；新建企业碳账户空户（如无已有账户）。",
        output_rows=None,
        ui_text="正式清单确认页，工具栏含【确认锁定】，表格含多选复选框及项目子明细展开行。",
        rule_lines=[
            "正式清单展示字段（元）：月均贷款余额（avgMonthlyBalance×10000）、平均资产总额（avgTotalAssets×10000）、年报营业收入（operatingRevenue×10000）。",
            "确认锁定后立即建档企业碳账户（creditCode为主键），用于后续归集排放记录。",
        ],
        data_text="status: draft→confirmed；lockedAt时间戳；碳账户以creditCode为主键建档。",
    )

    # ── 4.1.4 数据采集 ──────────────────────────────────────────────────────
    hp(doc, "4.1.4数据采集", 3)
    np(doc, "数据采集是投融资碳核算流程第四步。从v0.6起，数据采集增加「客户归集合并」层：同一客户的同类贷款合并为一个收集单元下发，以减少客户经理重复填报。具体归集规则见3.4节。")

    feat(doc,
        no="CR21093-A_ICC_GN0401",
        name="数据采集台账查询与筛选",
        desc="查询已锁定正式清单的数据采集进度（归集单元视图），支持多维度筛选。",
        input_rows=[
            ("客户名称", "文本框（模糊）", "否", "100", "", "", ""),
            ("核算方法", "下拉框", "否", "50", "", "", ""),
            ("收集状态", "下拉框", "否", "20", "全部", "全部/待处理/收集中/已完成/已退回", "5档"),
            ("审核状态", "下拉框", "否", "20", "", "全部/待审核/已通过/已退回", ""),
        ],
        flow_lines=[
            "进入数据采集页，展示统计卡片（已锁定笔数/已归集单元数/须收集单元数/已完成/已退回）。",
            "列表按归集单元（collectGroup）展示：客户名称、信用代码、业务种类（非项目/项目）、合并笔数、下发分行、主办客户经理、归集核算行业、采集状态、操作列。",
            "点击行展开查看该归集单元内的逐笔贷款明细（贷款账号、信贷品种、一级分行、余额等，只读）。",
            "操作列：已派发记录含【查看填报】；通过审核记录含【退回】（具备权限时展示）。",
        ],
        check_lines=["客户经理不可访问数据采集管理页面。"],
        output_text="数据采集列表（归集单元视图）；统计卡片。",
        output_rows=None,
        ui_text="数据采集页，工具栏含【发放收集任务】【调取接口数据】【一键提交数据】【信贷数据兜底法】。",
        rule_lines=[
            "收集状态5档：待处理（无排放数据）/收集中（已派发填报中）/已完成（已通过审核或已获取排放数据）/已退回。",
            "展示「系统核算方法」时须展示接口数据来源（如：格澜数据-各地区企业环境信息披露平台）。",
        ],
        data_text="collectGroup.status映射收集状态；归因排放展示需从calculations取值。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0402",
        name="调取接口数据（格澜报告法+经济法直算）",
        desc="一键从格澜等外部接口拉取报告法主体排放，并对符合条件的非项目归集单元执行经济活动法直算。",
        input_rows=[("任务ID", "系统参数", "是", "50", "", "", "")],
        flow_lines=[
            "用户点击【调取接口数据】。",
            "格澜接口调取：对已锁定且无主体排放、且非「项目（以项目方式计算）」的记录按信用代码调用格澜接口（一个归集单元调一次），写入gelanEntityEmission、gelanFetchedAt、gelanStatus=success。",
            "格澜成功记录：同步生成归因排放及calculation记录，归因排放写入对应企业碳账户annualProfiles；预填至客户经理填报页「其他来源报告法」Tab，数据来源标注「格澜数据-各地区企业环境信息披露平台」，客户经理可编辑。",
            "格澜无数据记录：标记gelanStatus=no_data，不写入排放数值。",
            "经济法直算：对collectMode=economy_direct、尚无主体排放的非项目归集单元执行：E_主体=年报营业收入(万元)×行业因子；E_归因=E_主体×月均余额/平均资产总额（见3.6节）。",
            "操作完成后展示：「格澜N笔已获取，M笔无数据；经济法直算K笔」。",
        ],
        check_lines=[
            "项目贷款（以项目方式计算）不调取格澜主体排放，不执行经济法直算，须通过发放收集任务由客户经理填报项目排放。",
            "若营收为0或缺失，经济法直算无法执行，提示具体原因。",
        ],
        output_text="格澜报告法排放数据；经济法直算排放数据；提示文案。",
        output_rows=None,
        ui_text="数据采集页工具栏【调取接口数据】按钮。",
        rule_lines=[
            "格澜接口数据预填至「其他来源报告法」Tab，格澜来源固定文案：「格澜数据-各地区企业环境信息披露平台」。",
            "经济法直算公式见3.6节；行业因子从因子库按gbIndustryCode匹配经济活动法因子。",
        ],
        data_text="gelanEntityEmission、gelanFetchedAt、gelanStatus；economyDirectStatus=done；calculation记录qualityScore=4。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0403",
        name="发放收集任务",
        desc="为已锁定正式清单记录（按归集单元）创建客户经理数据采集任务。",
        input_rows=[("勾选的归集单元ID列表", "复选框", "是", "—", "", "", "")],
        flow_lines=[
            "勾选归集单元（collectGroup），点击【发放收集任务】。",
            "系统按归集单元创建supplement：每个collectGroup对应1个supplement，写入下发分行、主办客户经理、归集成员贷款列表、归集核算行业、截止日期。",
            "派发后收集状态更新为「收集中」，客户经理在任务清单中可见。",
        ],
        check_lines=["已派发的归集单元不重复派发。"],
        output_text="收集任务（supplement）；提示派发单元数。",
        output_rows=None,
        ui_text="数据采集页行选择复选框和【发放收集任务】按钮。",
        rule_lines=[
            "已通过格澜/经济法获得排放底稿的归集单元，仍可发放收集任务给客户经理，数据采集与直算并行。",
            "收集任务中展示归集单元内所有关联贷款摘要（账号/分行/余额），供客户经理参考。",
        ],
        data_text="supplement.collectGroupId关联归集单元；supplement.memberFormalIds为逐笔正式清单ID列表。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0404",
        name="信贷数据兜底法",
        desc="当无法获取营收等主体数据时，以月均贷款余额×行业排放因子估算主体排放及归因排放。",
        input_rows=[("任务ID", "系统参数", "是", "50", "", "", "")],
        flow_lines=[
            "用户点击【信贷数据兜底法】。",
            "系统对无主体排放的正式清单记录计算：E_主体=月均余额(万元)×行业因子；E_归因=E_主体（全额归因，不按资产折算），见3.6节。",
            "写入calculation记录，source=credit_fallback，方法标记「其他计算法」；质量得分=5（最低）。",
        ],
        check_lines=["仅对尚无主体排放的记录应用，不覆盖已有格澜/经济法/手动数据。"],
        output_text="兜底法计算结果；提示处理笔数。",
        output_rows=None,
        ui_text="数据采集页工具栏【信贷数据兜底法】按钮。",
        rule_lines=[
            "信贷数据兜底法对应「其他计算法」，DQR质量等级=5；月均余额以万元为单位传入因子计算。",
        ],
        data_text="calculation.source=credit_fallback；qualityScore=5；entityEmission=avgLoanBalance×industryFactor；attributedEmission=entityEmission。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0405",
        name="一键提交数据",
        desc="当所有已锁定记录均有主体排放后，将任务推进到排放计算环节。",
        input_rows=[("任务ID", "系统参数", "是", "50", "", "", "")],
        flow_lines=[
            "点击【一键提交数据】。",
            "系统校验：所有confirmed状态的正式清单记录均已有主体排放数据。",
            "推进workflowStep至排放计算；跳转排放计算页。",
        ],
        check_lines=["存在缺失主体排放时按钮禁用，提示「请待全部记录计算出主体排放后再提交」。"],
        output_text="采集完成状态；跳转排放计算页。",
        output_rows=None,
        ui_text="数据采集页【一键提交数据】按钮。",
        rule_lines=["信贷数据兜底法可补填剩余缺失排放记录以满足提交条件。"],
        data_text="workflowStep→排放计算（4）；dataCollectSubmitted=true。",
    )

    # ── 4.1.5 数据采集（客户经理） ─────────────────────────────────────────
    hp(doc, "4.1.5数据采集（客户经理）", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0501",
        name="客户经理任务清单",
        desc="客户经理查看本人数据采集任务列表和状态。",
        input_rows=[
            ("任务ID", "系统参数", "是", "50", "", "", ""),
            ("当前用户", "系统参数", "是", "50", "", "", ""),
        ],
        flow_lines=[
            "进入客户经理任务清单页。",
            "系统按主办客户经理（manager字段）过滤supplement列表，展示本人任务。",
            "列表字段：客户名称、合并笔数、填报截止日期、状态、操作（【去填报】/【查看】）。",
        ],
        check_lines=["客户经理只见本人任务；不展示数据管理端功能入口。"],
        output_text="本人收集任务列表。",
        output_rows=None,
        ui_text="客户经理任务清单页；状态展示：待填报/填报中/已提交/已通过/已退回。",
        rule_lines=["状态说明：待填报（未填报）/填报中（暂存）/已提交（待审核）/已通过（审核通过）/已退回（需重填）。"],
        data_text="supplement.manager字段过滤；supplement.status映射状态文案。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0502",
        name="在线数据采集填报",
        desc="客户经理通过在线表单补充企业基本信息、多方法排放数据和佐证附件，并提交至审核流程。收集填报字段根据行业+方法模板动态展示。",
        input_rows=[
            ("企业基本信息", "表单", "否", "—", "", "", "包含营收、资产等数据核实"),
            ("报告法-政府/权威来源", "表单+附件", "否", "—", "", "", "须上传附件；字段见3.5.1"),
            ("报告法-其他来源（格澜预填）", "表单", "否", "—", "", "", "可编辑；字段见3.5.2"),
            ("物理活动法-能源法", "动态表单", "否", "—", "", "", "字段由方法模板决定，见3.5.3"),
            ("物理活动法-产品法", "动态表单", "否", "—", "", "", "字段由方法模板决定，见3.5.4"),
            ("经济活动法-营收法", "只读展示", "否", "—", "", "", "系统预填，不可编辑，见3.5.5"),
            ("项目信息（项目类业务）", "表单", "否", "—", "", "", "仅项目类展示"),
        ],
        flow_lines=[
            "客户经理点击【去填报】进入碳排放信息采集页。",
            "系统根据 归集单元核算行业 + 业务种类 查找对应方法模板（status=published）；按模板配置动态渲染各方法Tab字段。",
            "报告数据分为两Tab：「政府/权威来源报告法」（须附件）和「其他来源报告法」（格澜预填，可编辑）。",
            "能源法字段按模板展示：固定燃料输入行 + 其他燃料动态行（下拉+数量）+ 所属电网+净购入电量 + 过程排放动态行。",
            "产品法字段按模板展示：按产品类型分组的产量输入字段。",
            "经济活动法-营收法：系统预填年报营业收入和行业因子，只读展示计算结果，不可编辑。",
            "点击【暂存】保存中间状态（status=in_progress）；点击【提交数据】提交审核（status=completed）。",
        ],
        check_lines=[
            "处于待审核、已通过状态时不可编辑。",
            "经济活动法-营收法预填值只读，不可编辑。",
            "报告法-政府/权威来源Tab中，「该数据是否经政府/第三方核查=是」时须上传附件，否则阻断提交。",
            "附件：每个Tab最多3个，单文件不超过20MB，格式：pdf/doc/docx/xls/xlsx/png/jpeg/jpg。",
        ],
        output_text="填报数据（暂存或提交至审核）。",
        output_rows=None,
        ui_text="碳排放信息采集页，含「企业基本信息」「排放数据」「审批流程」三个主Tab；排放数据Tab内按方法分子Tab。底部含【暂存】【提交数据】按钮。",
        rule_lines=[
            "方法模板联动：系统先查找publishedTemplate(industryMajor, accountingType, methodId)；未找到时用defaultMode（仅显示排放总量输入框）。",
            "「其他来源报告法」格澜预填数据的来源文案固定为「格澜数据-各地区企业环境信息披露平台」；客户经理可修改数值但来源提示保留。",
            "系统提示文案统一使用「数据采集」，不使用「补录」；审核文案统一使用「退回」，不使用「驳回」。",
        ],
        data_text="fieldData分对象存储：{reportAuthority:{...}, reportOther:{...}, energy:{...}, product:{...}}；supplement.status=in_progress（暂存）或completed（提交）。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0503",
        name="项目信息填报",
        desc="项目类业务在收集填报中补充项目信息，并联动确定业务种类终态。",
        input_rows=[
            ("是否可提供项目信息", "下拉框", "否", "10", "", "是/否", "仅项目类业务展示"),
            ("项目信息相关字段", "表单", "条件必填", "—", "", "", "选择「是」时展开"),
        ],
        flow_lines=[
            "项目类业务进入填报页，「企业基本信息」Tab显示「是否可提供项目信息」字段。",
            "选「是」→ 展开项目信息填报区；提交后accountingType=project_as_project（终态）。",
            "选「否」→ 不展示项目字段；提交后accountingType=project_as_non_project（终态）。",
        ],
        check_lines=["仅项目类业务展示该字段（bizType=project）。"],
        output_text="projectDetails或不可提供标记；accountingType更新为终态。",
        output_rows=None,
        ui_text="企业基本信息Tab内联动展示项目信息填报区。",
        rule_lines=["若接口已同步projectDetails，则项目信息只读展示，不允许客户经理在填报页清空。"],
        data_text="projectInfoAvailable=true/false；projectDetails=[]或具体项目数组；accountingType终态定档。",
    )

    # ── 4.1.6 数据审核 ──────────────────────────────────────────────────────
    hp(doc, "4.1.6数据审核", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0601",
        name="查询审核记录列表",
        desc="总行、分行按角色查看收集数据的审核记录，支持多维筛选。",
        input_rows=[
            ("任务名称", "文本框（模糊）", "否", "100", "", "", ""),
            ("核算年度", "下拉框", "否", "4", "", "", ""),
            ("审核环节", "下拉框", "否", "20", "", "全部/分行初审/总行终审", ""),
            ("审核状态", "下拉框", "否", "20", "", "全部/待审核/通过/退回", ""),
        ],
        flow_lines=[
            "进入数据审核页，筛选面板含：任务名称（模糊）、核算年度、审核环节（独立字段）、审核状态（独立字段）；查询/重置按钮同行展示。",
            "系统按角色过滤审批记录：总行见总行环节；分行见辖内分行环节；客户经理仅见本人进度。",
            "列表字段：任务名称、核算年度、客户名称、审核环节（分行初审/总行终审）、审核状态（待审核/通过/退回，两字段独立展示）、提交人、提交时间、操作（查看/审核）。",
        ],
        check_lines=["客户经理不可执行审核操作，仅查看进度。"],
        output_text="审核列表。",
        output_rows=None,
        ui_text="数据审核列表页；审核环节与审核状态拆分为独立字段。",
        rule_lines=[
            "审核环节（分行初审/总行终审）与审核状态（待审核/通过/退回）为2个独立字段，不合并展示。",
            "总行发起任务：分行初审通过后进入总行终审；分行发起任务：分行审核通过即完成。",
        ],
        data_text="approval.reviewLevel=branch/hq；approvalStatus=pending/approved/rejected。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0602",
        name="审核详情与处理",
        desc="审核人查看收集填报内容并执行通过、退回操作；分行审核通过时须明确最终采用方法；支持调整归属行业及适用因子。",
        input_rows=[
            ("approvalId", "系统参数", "是", "50", "", "", ""),
            ("最终采用方法", "下拉框", "条件必填", "50", "", "", "分行审核通过时必选"),
            ("退回原因", "文本域", "条件必填", "500", "", "", "退回时必填"),
            ("归属行业调整", "四级级联选择", "否", "—", "", "", "审核时可调整"),
            ("适用排放因子调整", "因子选择", "否", "—", "", "", "审核时可从因子库选择"),
        ],
        flow_lines=[
            "审核人点击【审核】进入审核详情页，只读展示收集填报内容（含所有方法数据）和审批流程时间线。",
            "可在审核界面调整客户归属行业（国标行业分类）及对应适用排放因子（从因子库选择）；调整后系统重新计算该归集单元排放。",
            "若记录填写了多种方法数据，分行审核通过前须在弹窗中明确选择最终采用方法。",
            "点击【审核通过】→弹出方法选择确认弹窗（分行环节）→选定后更新审批记录，按任务类型流转至下一节点。",
            "点击【退回至客户经理】→填写退回原因（必填）→状态变更为「已退回」，客户经理可重新编辑提交。",
            "点击【本级修正】→分行绿金岗可在本级直接修改后重新提交，无需退回客户经理。",
        ],
        check_lines=[
            "退回时退回原因为必填，否则提示「请填写退回原因」。",
            "非当前审批角色不可审核。",
            "审核通过时若有多种方法数据，必须选定最终采用方法后方可通过。",
        ],
        output_text="审批结果；归属行业和因子调整记录。",
        output_rows=None,
        ui_text="审核详情页，含「填报内容」和「审批流程」两个Tab；操作区含【审核通过】【退回至客户经理】【本级修正】【取消】。",
        rule_lines=[
            "审核文案统一使用「退回」，不使用「驳回」（依据20260617会议纪要）。",
            "调整归属行业和因子后须记录审计日志；调整仅影响本次审核结果，不修改正式清单底层字段。",
        ],
        data_text="rejectReason；approvedMethodId；industryAdjustment；factorAdjustment；supplement.status=returned（退回）；auditStage记录当前审核层级。",
    )

    # ── 4.1.7 排放计算 ──────────────────────────────────────────────────────
    hp(doc, "4.1.7排放计算与结果确认", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0701",
        name="查询排放计算结果",
        desc="查询汇总主体排放、归因排放、DQR及分行业统计；所有已审核记录的归集单元主体排放已拆分为逐笔归因。",
        input_rows=[("任务ID", "系统参数", "是", "50", "", "", "")],
        flow_lines=[
            "进入碳排放计算页。",
            "展示统计卡片：总归因排放（tCO₂e）、已计算笔数、DQR数值和评级（A/B+/B/B-/C）。",
            "分行业归因排放图表（横向条形图）；下方逐笔计算结果列表（客户/行业/方法/主体排放/归因排放/质量等级/状态）。",
            "DQR计算：Σ(各笔归因排放 × qualityScore) / Σ(各笔归因排放)；质量等级映射见3.9节。",
        ],
        check_lines=["数据采集未完成时不允许进入排放计算页（workflowStep校验）。"],
        output_text="计算结果列表；统计卡片；DQR综合评级。",
        output_rows=None,
        ui_text="碳排放计算页，统计卡片行+行业图表+分页列表；底部含【确认结果】按钮。",
        rule_lines=[
            "逐笔归因排放公式见3.6节；各方法质量得分见3.9节。",
            "归集单元主体排放→逐笔归因拆分算法见3.8节。",
        ],
        data_text="calculation.entityEmission（主体排放）；attributedEmission（归因排放）；qualityScore（质量得分）；DQR综合值；qualityGrade（A/B+/B/B-/C）。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0702",
        name="确认结果并生成报告",
        desc="确认排放计算结果，将排放记录归集至企业碳账户，并生成报告导出记录。",
        input_rows=[
            ("任务ID", "系统参数", "是", "50", "", "", ""),
            ("导出范围", "下拉框", "是", "50", "", "监管报送范围/管理全量范围", ""),
            ("报告模板", "下拉框", "是", "50", "", "", ""),
            ("导出格式", "下拉框", "是", "20", "", "Word/Excel", ""),
        ],
        flow_lines=[
            "在排放计算页点击【确认结果】，系统写入resultsConfirmed=true。",
            "逐笔calculation记录归集至企业碳账户annualProfiles（按creditCode找到账户，写入该年度排放数据）。",
            "跳转至生成报告页，用户选择报告范围、模板和格式后点击导出，生成报告记录展示在历史列表。",
        ],
        check_lines=["未确认结果不可导出正式报告。"],
        output_text="确认状态；报告导出记录；企业碳账户归集排放数据。",
        output_rows=None,
        ui_text="排放计算页【确认结果】按钮；生成报告页含报告配置和历史导出列表。",
        rule_lines=[
            "企业碳账户归集逻辑：对每笔formal记录，找到creditCode对应账户，写入annualProfiles[year]={entityEmission, attributedEmission, method, ...}。",
            "同一账户同一年度多笔记录：entityEmission取最新确认值；attributedEmission累加。",
        ],
        data_text="resultsConfirmed；annualProfiles[year]写入排放摘要；carbonAccountRecords存储逐笔排放记录快照。",
    )

    # ── 4.1.8 台账管理 ──────────────────────────────────────────────────────
    hp(doc, "4.1.8台账管理", 3)

    feat(doc,
        no="CR21093-A_ICC_GN0801",
        name="查询投融资碳核算台账",
        desc="按分行、核算年度、客户等条件查询已纳入核算的业务台账，并支持导出。",
        input_rows=[
            ("一级分行", "下拉框", "否", "50", "全部", "", ""),
            ("核算年度", "下拉框", "否", "4", "", "", ""),
            ("客户名称", "文本框（模糊）", "否", "100", "", "", ""),
        ],
        flow_lines=[
            "进入台账管理页，展示筛选面板（一级分行/核算年度/客户名称/查询/重置）。",
            "列表展示：序号、任务名称、核算年度、投向行业范围、所属行业范围。",
            "点击行或【查看】跳转至该任务的排放计算清单详情页。",
            "工具栏【导出台账】，导出当前筛选结果。",
        ],
        check_lines=["客户经理不展示台账管理菜单。"],
        output_text="台账列表；导出文件。",
        output_rows=None,
        ui_text="台账管理列表页，顶部筛选面板，下方分页表格，工具栏含【导出台账】。",
        rule_lines=["总行可查看全行台账；分行仅查看辖内分行台账。"],
        data_text="从formalList关联tasks、calculations等实体；exportLedger生成导出文件。",
    )

    # ── 4.1.9 排放因子库 ────────────────────────────────────────────────────
    hp(doc, "4.1.9排放因子库", 3)
    np(doc, "排放因子库存储核算所需的各类排放因子，分内置因子（isBuiltin=true，只读）和自定义因子（isBuiltin=false，可维护）两类。不再使用版本年度概念，以口径标签（caliberTag=pbo/bank）区分人行口径与行自定义口径。因子库数据通过「新增」或「导入」维护，由数据采集环节（经济法直算/兜底法）和审核环节（因子调整）引用。")

    feat(doc,
        no="CR21093-A_ICC_GN0901",
        name="查询排放因子",
        desc="查询指引内置因子和自定义因子，支持按方法、行业、口径标签多维筛选。",
        input_rows=[
            ("计算方法", "下拉框", "否", "50", "", "报告法/能源法/产品法/经济活动法/其他", ""),
            ("行业大类", "下拉框", "否", "50", "", "全量行业大类", ""),
            ("口径标签", "下拉框", "否", "20", "", "人行口径/行自定义", ""),
            ("关键词", "文本框", "否", "100", "", "", ""),
        ],
        flow_lines=[
            "进入排放因子库，展示统计卡片（因子组数：全部/能源法/产品法/经济活动法/自定义）。",
            "工具栏含【新增因子】【导入因子】（跳转至独立导入页）。不展示「复制因子版本」和「版本年度」。",
            "按筛选条件过滤，表格展示：计算方法、行业、名称/细分项、因子值、单位、口径、来源、操作列。",
            "内置因子操作：【复制为自定义】【查看】；自定义因子操作：【编辑】【删除】。",
        ],
        check_lines=["内置因子不可直接编辑或删除。"],
        output_text="因子列表；统计卡片。",
        output_rows=None,
        ui_text="排放因子库列表页；无「版本年度」列和「复制因子版本」按钮。",
        rule_lines=[
            "产品法因子分两套口径：人行口径（caliberTag=pbo）和行自定义（caliberTag=bank）。",
            "因子匹配优先级：自定义因子优先于内置因子；GB国标四级代码匹配优先于行业大类匹配。",
        ],
        data_text="isBuiltin=true只读；isBuiltin=false可维护；无versionYear字段；caliberTag=pbo/bank。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0902",
        name="新增与编辑排放因子",
        desc="新增自定义因子或编辑已有自定义因子，选择方法后联动展示专属字段。",
        input_rows=[
            ("计算方法", "下拉框", "是", "50", "", "能源法/产品法/经济活动法/其他", ""),
            ("行业（GB/T 4754四级）", "级联下拉", "是", "—", "", "", ""),
            ("名称/细分项", "文本框", "是", "100", "", "", ""),
            ("因子值", "数字输入", "是", "20,6", "", ">0", ""),
            ("单位", "文本框", "是", "50", "", "tCO₂e/万元营收等", ""),
            ("口径标签", "下拉框", "否", "20", "", "人行口径/行自定义", ""),
            ("来源说明", "文本域", "是", "500", "", "", ""),
        ],
        flow_lines=[
            "点击【新增因子】或编辑自定义因子进入表单页。",
            "选择计算方法后联动展示专属字段：能源法→排放源类型/细分项/子行业；产品法→主要产品/细分项/口径标签；经济活动法→GB国标行业选择。",
            "点击【保存】写入因子库（isBuiltin=false）。",
            "内置因子可点击【复制为自定义】，生成带来源说明的副本并进入编辑页。",
        ],
        check_lines=["来源说明为必填。", "内置因子不可直接编辑。"],
        output_text="自定义因子记录。",
        output_rows=None,
        ui_text="新增/编辑排放因子页，两栏表单，底部含【保存】【取消】。",
        rule_lines=["【待确认】因子发布审批流程和正式生效规则。"],
        data_text="自定义因子ID前缀EF-C；isBuiltin=false；caliberTag可选。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN0903",
        name="导入排放因子（独立导入页）",
        desc="通过独立导入页，以两步流程批量导入排放因子，支持查看导入历史。",
        input_rows=[("因子导入文件", "文件上传", "是", "—", "", "Excel/CSV", "")],
        flow_lines=[
            "点击【导入因子】进入独立导入页（#/factors/import）。",
            "Step1：页面提供【下载导入模板】按钮，用户下载标准格式模板。",
            "Step2：用户上传文件，系统解析校验格式，预览识别到的条目数。",
            "点击【确认导入】：写入符合格式的因子记录（isBuiltin=false），重复factorGroupKey自动跳过。",
            "导入完成后展示：「已添加N条，已跳过M条（重复）」。",
            "页面底部展示导入历史列表（时间/文件名/添加数/跳过数/状态）。",
        ],
        check_lines=[
            "导入文件字段须与模板一致；格式错误行自动跳过并在日志中说明。",
            "重复因子（按methodId+industryMajor等效factorGroupKey判断）跳过，不覆盖。",
        ],
        output_text="批量导入的自定义因子；导入历史记录。",
        output_rows=None,
        ui_text="独立导入页（#/factors/import），含两步流程说明、文件选择区域、预览、历史导入列表。",
        rule_lines=["生产环境应持久化导入历史到服务端。"],
        data_text="批量写入factors，isBuiltin=false；factorGroupKey去重；importHistory存储批次结果。",
    )

    # ── 4.1.10 企业碳账户 ───────────────────────────────────────────────────
    hp(doc, "4.1.10企业碳账户", 3)
    np(doc, "企业碳账户以统一社会信用代码为主键（同一主体一个账户），核算任务确认结果后自动归集排放记录。账户按企业主体持续累积，不因年度任务重置而清库。")

    feat(doc,
        no="CR21093-A_ICC_GN1001",
        name="查询企业碳账户",
        desc="按核算年度查询已归集形成的企业碳账户，展示主体排放汇总和账户状态。",
        input_rows=[
            ("核算年度", "年度切换标签", "否", "4", "", "", ""),
            ("关键词", "文本框", "否", "100", "", "", "企业名称/信用代码模糊"),
            ("分行", "下拉框", "否", "50", "", "", ""),
            ("账户状态", "下拉框", "否", "20", "", "全部/正常/停用/注销", ""),
        ],
        flow_lines=[
            "进入企业碳账户菜单，按角色和组织权限过滤账户（总行全行；分行辖内）。",
            "统计卡片：正常账户数、主体排放合计（tCO₂e）、含项目子账户数。",
            "列表：序号、企业名称（含项目子账户缩进标注）、统一社会信用代码、客户号、核算方法、主体排放（tCO₂e）、账户状态、操作（查看/编辑/状态变更）。",
        ],
        check_lines=["客户经理不展示企业碳账户菜单。"],
        output_text="账户列表；统计卡片。",
        output_rows=None,
        ui_text="企业碳账户列表页，年度切换标签栏+统计卡片+筛选区+分页表格；项目子账户以缩进子行展示。",
        rule_lines=[
            "账户主键=统一社会信用代码（creditCode）；同一主体多条贷款→同一账户。",
            "账户允许空户（已建档但当年无确认数据）。",
        ],
        data_text="account.status=active/suspended/closed；annualProfiles存各年度排放摘要；carbonAccountRecords存逐笔记录。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN1002",
        name="查看/编辑碳账户详情",
        desc="查看账户档案、排放明细和年度排放趋势；active状态账户可进入编辑模式修改档案。",
        input_rows=[
            ("accountId", "系统参数", "是", "50", "", "", ""),
            ("核算年度", "年度切换", "否", "4", "", "", "仅账户档案Tab显示"),
        ],
        flow_lines=[
            "点击账户详情，展示两个Tab：账户档案（含年度切换器）、趋势分析。",
            "账户档案Tab：展示企业名称、信用代码、行业、核算方法、主体排放（tCO₂e）等档案字段。active状态账户可进入编辑模式修改档案（不改变正式清单和核算计算结果）。",
            "查看模式：所有字段只读，但各方法排放Tab可切换查看（只读）。",
            "趋势分析Tab：折线图1=年度主体排放趋势（含数值标注，点位与年份严格对齐）；折线图2=碳强度趋势（主体排放÷营业收入，单位tCO₂e/万元营收；营收=0时虚线点显示「—」）；下方年度明细表（年度/笔数/主体排放/归因排放/营业收入/碳强度）。",
            "总行角色可执行账户状态变更（正常→停用→注销），操作记入操作日志。",
        ],
        check_lines=[
            "停用/注销账户仅可查看，不可编辑档案。",
            "分行不可变更账户状态。",
        ],
        output_text="账户详情；趋势图表；年度明细。",
        output_rows=None,
        ui_text="碳账户详情页，两Tab布局；「核算年度」切换器仅在账户档案Tab显示。趋势图所有数据点均显示数值标注。",
        rule_lines=[
            "碳强度=主体排放(tCO₂e) / 年报营业收入(万元)，单位tCO₂e/万元营收；营收来源annualProfiles[year].operatingRevenue。",
            "趋势图X轴严格对齐核算年度，不插值；无数据年份留空（不连线）。",
        ],
        data_text="annualProfiles[year]存档案摘要；trendByYear(records)汇总各年度数据；碳强度=entity/revenue。",
    )

    # ── 4.1.11 基础配置 ─────────────────────────────────────────────────────
    hp(doc, "4.1.11基础配置", 3)
    np(doc, "基础配置是系统管理功能集合，入口位于左侧侧栏，默认不显示（通过权限管理控制）。包含：参数字段库、方法模板、行业配置三个子模块。这三个子模块共同支撑数据采集收集填报页的字段动态渲染，开发须理解联动逻辑（见3.7节）。")

    feat(doc,
        no="CR21093-A_ICC_GN1101",
        name="参数字段库",
        desc="维护方法模板中可复用的参数字段定义，包括字段名称、数据类型、单位、格式等配置，供方法模板引用。",
        input_rows=[
            ("字段名称", "文本框", "是", "100", "", "", ""),
            ("数据类型", "下拉框", "是", "20", "", "数值型（number）/选项型（option）", ""),
            ("格式/精度", "文本框", "否", "50", "", "", "如：小数点后4位"),
            ("单位", "文本框", "否", "50", "", "", "如：万吨/MWh"),
            ("枚举选项", "文本域", "条件必填", "—", "", "", "选项型时填写，每行一个选项"),
        ],
        flow_lines=[
            "进入参数字段库页，列表展示已配置的参数字段。",
            "支持新增、编辑、删除参数字段定义。",
            "参数字段被方法模板引用后，方法模板编辑时可直接选取，避免重复定义。",
        ],
        check_lines=["参数字段删除前须检查是否被方法模板引用；已引用时不可删除。"],
        output_text="参数字段列表。",
        output_rows=None,
        ui_text="参数字段库列表页，工具栏含【新增参数】，下方为分页表格。",
        rule_lines=[
            "数值型字段：存储为number类型，可配置小数精度和单位。",
            "选项型字段：存储为string类型，枚举选项列表在方法模板中渲染为下拉框。",
        ],
        data_text="paramField实体：id、name、format（number/option）、unit、precision、options[]；被template.methods.*.fuelFixed等引用。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN1102",
        name="方法模板配置",
        desc="按行业+业务类型+核算方法维护采集字段模板，支持字段选择、公式绑定、因子绑定；发布后决定对应场景的收集填报页字段。",
        input_rows=[
            ("行业", "下拉框", "是", "50", "", "全量行业分类", ""),
            ("业务类型", "下拉框", "是", "50", "", "非项目/项目", ""),
            ("核算方法", "下拉框", "是", "50", "", "报告法/能源法/产品法/经济活动法", ""),
            ("采集字段列表", "多选/拖拽", "是", "—", "", "从参数字段库选取", ""),
            ("计算公式", "代码域", "否", "—", "", "", ""),
            ("关联因子", "因子选择", "否", "—", "", "从因子库选取", ""),
        ],
        flow_lines=[
            "进入方法模板配置页，列表展示已有模板（按行业+方法分组）。",
            "点击【新增模板】或编辑，进入四步向导：Step1基础信息→Step2采集字段（从参数字段库选取并配置是否必填/默认值）→Step3输入公式（填写计算公式说明）→Step4因子绑定（指定使用哪个排放因子）。",
            "点击【保存草稿】：status=draft，不影响现有填报页；点击【发布】：status=published，立即生效。",
            "模板发布后，对应行业+方法组合的填报页将使用该模板的字段渲染（见3.7节联动机制）。",
        ],
        check_lines=[],
        output_text="方法模板记录（draft/published）。",
        output_rows=None,
        ui_text="方法模板列表页+四步配置向导。",
        rule_lines=[
            "模板发布后对应行业+方法的填报页立即采用新模板；草稿状态不影响现有填报。",
            "若某行业+方法无已发布模板，填报页回退到defaultMode（仅展示排放总量输入框）。",
            "【待确认】模板发布是否需要审批流程（需业务确认）。",
        ],
        data_text="template.status=draft/published；template.methods.energy/product/report包含字段配置；resolveTemplate(supplement)在填报时查找对应已发布模板。",
    )

    feat(doc,
        no="CR21093-A_ICC_GN1103",
        name="行业配置",
        desc="维护国民经济行业分类（GB/T4754）四级结构，标识「人行八大高碳」「我行主要行业」；标识数据用于核算任务行业范围选项。",
        input_rows=[
            ("行业类别代码", "文本框", "是", "10", "", "GB/T4754四级代码", ""),
            ("行业名称", "文本框", "是", "100", "", "", ""),
            ("级别", "下拉框", "是", "10", "", "一/二/三/四级", ""),
            ("标识", "多选框", "否", "—", "", "人行八大高碳/我行主要行业", ""),
        ],
        flow_lines=[
            "进入行业配置页，若未导入数据展示提示。",
            "点击【导入全部国民经济行业分类】：加载GB/T4754-2017全量四级行业（含人行八大高碳和我行主要行业预置标识）。",
            "列表展示4列（一/二/三/四级行业名称及代码）+1列标识（人行八大高碳/我行主要行业/空）。",
            "支持逐行编辑标识；支持新增自定义行业条目；支持删除非内置条目。",
        ],
        check_lines=[],
        output_text="行业分类列表；标识配置。",
        output_rows=None,
        ui_text="行业配置页，工具栏含【导入全部国民经济行业分类】【新增行业】；下方分页表格（4列分类+1列标识）。",
        rule_lines=[
            "「人行八大高碳」标识决定核算任务行业范围选项「人行投融资碳核算八大高碳行业」的实际行业代码集合。",
            "「我行主要行业」标识决定「人行八大高碳+我行主要行业」选项新增的行业集合。",
        ],
        data_text="industryConfig.rows存储行业记录；tag枚举：pbo（人行八大高碳）/bank（我行主要行业）；IndustryScope.getEightCodes()从标识pbo的记录取值。",
    )

    # ── 4.1.12 权限管理 ─────────────────────────────────────────────────────
    hp(doc, "4.1.12权限管理", 3)

    feat(doc,
        no="CR21093-A_ICC_GN1201",
        name="菜单可见性配置",
        desc="通过权限管理页控制左侧菜单项展示，默认基础配置模块不展示。",
        input_rows=[
            ("菜单ID", "系统字段", "是", "50", "", "", ""),
            ("是否显示", "复选框", "是", "—", "", "勾选/取消", ""),
        ],
        flow_lines=[
            "总行角色进入权限管理页，列出全部可配置菜单项和复选框。",
            "默认：核算任务管理/数据审核/因子库/接口管理/碳账户/台账管理等主功能菜单可见；基础配置（参数字段库/方法模板/行业配置）默认不可见。",
            "修改后点击【保存配置】，侧栏菜单实时刷新。",
        ],
        check_lines=[],
        output_text="菜单可见性配置；侧栏实时刷新。",
        output_rows=None,
        ui_text="权限管理页，菜单树+复选框；底部含【保存配置】。",
        rule_lines=["正式生产权限配置应沉淀到服务端角色权限体系，不依赖前端localStorage。"],
        data_text="MenuPermissions模块维护visibilityMap；isVisible(menuId,roleKey)供菜单渲染调用。",
    )

    # ── 4.2 外部接口需求 ─────────────────────────────────────────────────────
    hp(doc, "4.2外部接口需求", 2)
    add_plain_table(doc, ["外部接口", "说明", "待确认事项"], [
        ["信贷台账接口", "按月批次推送贷款台账；字段含分行、客户、品种、月均余额、月末余额合计、华夏存续月份、营收、资产（当年末+上年末）、行业、主办客户经理、项目信息（projectDetails）等。", "正式接口字段名/类型；调用频率；错误码；重试；脱敏要求。"],
        ["格澜企业排放接口", "按信用代码查询企业主体碳排放（报告法数据）；返回ghgTotalEmission、scope1/2Emission、reportYear、reportSource等字段。", "生产接口地址；认证方式；请求参数；返回字段；频率与缓存策略。"],
        ["审批与待办接口", "提交审核传入单据类型/业务ID/提交人；接收审批状态回写（通过/退回/作废）。", "正式审批状态码；流程实例ID；退回后解锁字段范围。"],
        ["组织权限接口", "复用绿金系统组织树、用户、角色和数据权限。", "角色枚举和权限控制粒度。"],
        ["附件与文件服务", "收集填报须支持附件上传；病毒扫描、预览、留存年限。", "文件服务接口、大小/数量限制。"],
    ])

    # ── 4.3 待确认事项汇总（v0.6新增完整清单） ──────────────────────────────
    hp(doc, "4.3待确认事项汇总", 2)
    add_plain_table(doc, ["序号", "主题", "待确认内容", "影响范围"], [
        ["C01", "信贷品种→项目类映射",       "项目贷款/一般性固定资产贷款/出口退税账户托管贷款等完整清单及判断逻辑", "业务种类识别"],
        ["C02", "集团主办行字段",             "非项目类跨分行时，集团主办行（groupLeadBranch）字段从信贷接口/CRM/人工指定哪个来源取", "归集合并下发规则"],
        ["C03", "项目贷款主办分行字段",       "projectLeadBranch字段来源及多项目跨分行时的优先级规则（金额最大/首笔/接口指定）", "归集合并下发规则"],
        ["C04", "非项目单投向核算行业",       "非项目归集单元仅1个贷款投向行业时，核算行业用投向行业还是客户所属行业", "归集核算行业"],
        ["C05", "业务种类名称确认",           "三档终态名称是否与人行附件4完全一致（含是否带「贷款」二字）", "业务种类识别/展示"],
        ["C06", "营收/资产客户级字段规则",    "归集单元内多笔贷款，客户级年报营收、合并报表资产总额以哪笔/哪个系统为准（建议合并报表数仓取单次）", "归因排放计算"],
        ["C07", "方法模板发布审批",           "方法模板发布是否需要审批流程，由谁审批", "基础配置/方法模板"],
        ["C08", "格澜接口调取粒度",           "格澜接口是按正式清单笔调取还是按归集单元（creditCode）调取一次", "调取接口数据"],
        ["C09", "报告模板提供",               "监管与管理报告模板、导出维度由行方提供的清单与时间表", "生成报告"],
        ["C10", "碳账户排放结果更正",         "核算结果更正后，碳账户内对应记录是覆盖、冲正还是保留多版本", "企业碳账户"],
        ["C11", "碳账户停用规则",             "停用/注销账户是否仍参与历史统计；注销后是否允许再启用", "企业碳账户状态管理"],
        ["C12", "华夏银行重点行业清单",       "「人行八大高碳+我行主要行业」中的我行重点行业具体代码清单", "核算任务行业范围"],
        ["C13", "分行发起任务审批节点",       "分行发起任务时，分行审核通过是否还需要总行复核，还是直接进入排放计算", "数据审核流程"],
        ["C14", "数据采集归集合并是否分期",   "客户归集合并功能是否本期全量实施，还是分P0/P1/P2/P3分期实施（见改造计划文档）", "数据采集模块"],
    ])

    # ── 4.4 非功能性需求 ─────────────────────────────────────────────────────
    hp(doc, "4.4非功能性需求", 2)

    hp(doc, "4.4.1性能", 3)
    for item in [
        "系统应支持年度全行信贷台账量级的候选清单筛选、正式清单生成、数据采集跟踪和批量排放计算。",
        "数据层读写采用内存缓存，避免每次操作触发全量JSON序列化；数据迁移只执行一次。",
        "页面路由复用已挂载的SPA Shell，切换页面只刷新内容区域，不重建整页DOM。",
        "「调取接口数据」操作批量处理，单次读取完整数据后用Map索引进行O(1)查找。",
        "数据采集列表渲染预建supplement/candidate/calculation索引，避免每行多次遍历。",
        "【待确认】生产环境单任务最大台账笔数、并发用户数、报告导出最大记录数等性能指标待确认。",
    ]: lp(doc, item)

    hp(doc, "4.4.2安全性", 3)
    for item in [
        "三类角色（总行绿金部/分行负责人/客户经理）按宿主系统权限控制菜单和数据可见范围。",
        "总行绿金部：全行任务/审核记录/碳账户可见；可执行碳账户状态管理。",
        "分行负责人：仅辖内数据；可执行分行审核，不可变更碳账户状态。",
        "客户经理：仅访问客户经理任务清单和收集填报页；不可访问管理端功能。",
        "基础配置（参数字段库/方法模板/行业配置）默认不可见，通过权限管理配置控制。",
    ]: lp(doc, item)

    hp(doc, "4.4.3可靠性", 3)
    for item in [
        "接口失败支持重试；收集任务关键状态（暂存/提交/审核/退回/通过）可追溯。",
        "已通过收集数据被管理员退回时，须作废原审核记录并保留退回原因。",
        "格澜接口调取失败不阻断主流程，标记gelanStatus并允许重试或转其他路径。",
        "企业碳账户仅接收已确认结果的数据，保证账户内数据与任务确认结果一致。",
        "localStorage存储不超过4MB（含碳账户数据压缩）；配额满时自动压缩并提示。",
    ]: lp(doc, item)

    hp(doc, "4.4.4业务数据处理", 3)
    for item in [
        "月均贷款余额：系统从接口台账按12个月月末余额计算；部分月份有余额时按实有存续月份数计算均值，而非固定除以12（详见3.2节）。",
        "平均资产总额：上年末+当年末合并报表资产总额之和除以2（详见3.2.3节）；候选清单展示当年末资产，正式清单展示平均资产总额。",
        "行业识别优先级：候选清单筛选和业务种类识别均以贷款投向所属行业（gbIndustryCode）为主要依据；企业所属行业（industryMajor）作为辅助展示字段保留。",
        "客户归集合并：数据采集环节按3.4节规则归集；台账逐笔数据保留；归因排放仍按笔计算。",
        "方法优先级：报告法 > 物理活动法-能源法 > 物理活动法-产品法 > 经济活动法 > 其他计算法；分行审核通过时须明确选定最终采用方法。",
        "数据采集与直算并行：即使格澜/经济法已获取排放底稿，仍须完整保留派发收集任务的业务动作。",
        "格澜接口数据规则：成功返回后写入gelanEntityEmission并同步到企业碳账户；预填至「其他来源报告法」Tab，来源文案固定为「格澜数据-各地区企业环境信息披露平台」，客户经理可编辑数值。",
    ]: lp(doc, item)

    hp(doc, "4.4.5设计约束", 3)
    for item in [
        "页面风格、菜单框架、权限控制、审批入口符合绿金系统现有规范。",
        "审核文案统一使用「退回」，不使用「驳回」；数据采集文案统一使用「收集」，不使用「补录」。",
        "基础配置模块（参数字段库/方法模板/行业配置）默认不在侧栏展示，通过权限管理页面开启。",
        "能源法、产品法等计算估算公式以业务确认版本为准，当前原型暂不自动试算。",
    ]: lp(doc, item)

    # ── 4.5 数据管理目标 ─────────────────────────────────────────────────────
    hp(doc, "4.5数据管理目标", 2)
    add_plain_table(doc, ["数据对象", "管理要求"], [
        ["核算任务",       "保留任务配置（年度/行业/组织范围/余额口径）、流程步骤、发起组织、确认结果时间。"],
        ["候选/正式清单", "保留台账字段（含月均余额、存续月份、平均资产）、纳入状态、业务种类、项目明细、锁定状态、归集单元关联。"],
        ["收集任务",       "保留归集单元成员列表、填报字段（含各方法数据）、附件元数据、提交状态、审核状态、退回原因、格澜预填数据。"],
        ["计算结果",       "保留方法、主体排放、归因排放、质量得分/等级、计算完成时间；格澜/直算/兜底来源标记；DQR综合值。"],
        ["报告记录",       "保留导出范围、模板、格式、记录数、总排放、生成时间和操作人。"],
        ["企业碳账户",     "账户档案、年度排放摘要（annualProfiles）、逐笔排放记录（carbonAccountRecords）、状态变更日志；编辑档案不覆盖核算结果。"],
        ["因子库",         "内置/自定义因子分离；来源说明留存；导入历史留痕；无版本年度概念。"],
        ["方法模板",       "发布状态（draft/published）变更记录；引用字段版本快照。"],
        ["行业配置",       "全量行业分类数据及标识变更记录；标识变更影响后续任务范围配置。"],
    ])
    pending(doc, "RTO/RPO指标、备份周期、日志保留周期按行内灾备规范确认。")

    np(doc, "")
    np(doc, "— 文档结束 —")
    np(doc, "版本 v0.61 · 2026.06.29 · 起草：Cursor · 待评审")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_doc()
