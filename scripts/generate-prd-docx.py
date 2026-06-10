#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将项目需求说明书 Markdown 转为 Word"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'
PRD_VERSION = 'v0.1'
PRD_DATE = '20260527'
MD_FILE = DOCS / f'华夏银行绿金系统-投融资碳核算与企业碳账户-项目需求说明书-{PRD_DATE}-{PRD_VERSION}.md'
OUT_DOCX = DOCS / f'华夏银行绿金系统-投融资碳核算与企业碳账户-项目需求说明书-{PRD_DATE}-{PRD_VERSION}.docx'

MERMAID_HINTS = {
    'flowchart TB': (
        '【示意图】两大模块关系：投融资碳核算（任务→清单→补数→计算→确认结果/报告）'
        ' → 确认结果后归集 → 企业碳账户（账户→排放记录→查询/汇总/趋势）；'
        '账户主键为统一社会信用代码 + 贷款号。'
    ),
    'flowchart LR': (
        '【示意图】核算主流程：发起任务 → 清单识别 → 正式清单与边界 → 补数与审核 → '
        '排放计算 → 确认结果 → 报告输出；确认结果同时归集至碳账户。'
    ),
    'erDiagram': (
        '【示意图】核心数据对象：核算任务包含正式清单；正式清单关联补录与计算；'
        '碳账户包含排放记录；计算结果挂载至排放记录（关联任务标识）。'
    ),
}


def set_doc_font(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def strip_md_bold(text):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)


def add_rich_paragraph(doc, text, style=None):
    """段落内保留 **粗体**"""
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def parse_table_row(line):
    line = line.strip()
    if not line.startswith('|'):
        return None
    cells = [c.strip() for c in line.strip('|').split('|')]
    return cells


def is_separator_row(cells):
    if not cells:
        return False
    return all(re.match(r'^:?-+:?$', c.replace(' ', '')) for c in cells)


def add_table(doc, rows):
    if len(rows) < 1:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ''
            table.rows[ri].cells[ci].text = strip_md_bold(val)


def mermaid_to_hint(block):
    for key, hint in MERMAID_HINTS.items():
        if key in block:
            return hint
    return '【示意图】详见 Markdown 电子版中的流程/结构图。'


def convert(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    set_doc_font(doc)

    doc.core_properties.title = '华夏银行绿金系统-投融资碳核算与企业碳账户-项目需求说明书'
    doc.core_properties.subject = f'项目需求说明书 {PRD_VERSION}'

    title_lines = []
    i = 0
    while i < len(lines) and lines[i].startswith('# '):
        title_lines.append(lines[i][2:].strip())
        i += 1
    if title_lines:
        for ti, t in enumerate(title_lines):
            if ti == 0:
                p = doc.add_heading(t, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_heading(t, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = []
    in_mermaid = False
    mermaid_buf = []
    table_buf = []
    pending_list = None  # 'bullet' | 'number' | None

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        rows = []
        for ln in table_buf:
            cells = parse_table_row(ln)
            if cells and not is_separator_row(cells):
                rows.append(cells)
        add_table(doc, rows)
        table_buf = []

    def flush_mermaid():
        nonlocal in_mermaid, mermaid_buf
        if not in_mermaid:
            return
        block = '\n'.join(mermaid_buf)
        add_rich_paragraph(doc, mermaid_to_hint(block))
        in_mermaid = False
        mermaid_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```mermaid'):
            flush_table()
            in_mermaid = True
            mermaid_buf = []
            i += 1
            continue
        if in_mermaid:
            if stripped.startswith('```'):
                flush_mermaid()
            else:
                mermaid_buf.append(line)
            i += 1
            continue

        if stripped.startswith('|'):
            flush_mermaid()
            pending_list = None
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if stripped == '---':
            pending_list = None
            doc.add_page_break()
            i += 1
            continue

        if stripped.startswith('## '):
            pending_list = None
            doc.add_heading(strip_md_bold(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith('### '):
            pending_list = None
            doc.add_heading(strip_md_bold(stripped[4:]), level=2)
            i += 1
            continue
        if stripped.startswith('#### '):
            pending_list = None
            doc.add_heading(strip_md_bold(stripped[5:]), level=3)
            i += 1
            continue

        if re.match(r'^\d+\.\s+', stripped):
            pending_list = 'number'
            m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
            doc.add_paragraph(f'{m.group(1)}. {strip_md_bold(m.group(2))}', style='List Number')
            i += 1
            continue

        if stripped.startswith('- '):
            pending_list = 'bullet'
            add_rich_paragraph(doc, stripped[2:], style='List Bullet')
            i += 1
            continue

        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            add_rich_paragraph(doc, stripped.strip('*'))
            i += 1
            continue

        if not stripped:
            pending_list = None
            i += 1
            continue

        # skip duplicate title # lines at body
        if stripped.startswith('# '):
            i += 1
            continue

        pending_list = None
        add_rich_paragraph(doc, stripped)
        i += 1

    flush_table()
    flush_mermaid()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print('已生成:', out_path)


if __name__ == '__main__':
    if not MD_FILE.exists():
        raise SystemExit(f'找不到 Markdown 源文件: {MD_FILE}')
    convert(MD_FILE, OUT_DOCX)
