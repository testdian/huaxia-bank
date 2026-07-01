#!/usr/bin/env python3
"""基于 v1.0-20260629 需规生成 v1.0-20260630，写入 2026-06-30 原型修订内容。"""
from __future__ import annotations

import shutil
from pathlib import Path

try:
    from docx import Document
except ImportError:
    import subprocess
    import sys

    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / '华夏银行绿金系统-投融资碳核算+企业碳账户模块_需求规格说明书-v1.0-20260629.docx'
DST = ROOT / '华夏银行绿金系统-投融资碳核算+企业碳账户模块_需求规格说明书-v1.0-20260630.docx'

REVISION_DESC = (
    '根据2026年6月29日原型评审意见修订：'
    '①数据审核（退回选对象、保存/修改、因子选用默认值、方法排放预览）；'
    '②排放因子库（内置因子可维护、来源表号规范、新增因子含我行主要行业）；'
    '③能源法/产品法（附2全量字段、非必填、手动新增、除报告法外无附件）；'
    '④新建任务日期校验；'
    '⑤审批流程四步展示。'
)

PARA_REPLACEMENTS = {
    '支持分行初审时明确最终采用的核算方法，支持退回至客户经理；审核时可调整归属行业及适用因子。': (
        '支持分行初审时明确最终采用的核算方法（各方法展示主体排放预览或「缺少数据」提示）；'
        '支持退回至客户经理（须选择退回对象并填写原因）；审核时可调整归属行业、选用因子及因子数值。'
    ),
    '能源法字段分为「固定燃料」「其他燃料（动态行）」「净购入电量/热力」「过程排放（动态行）」四个部分，具体字段由对应行业的方法模板配置决定。以下列出系统内已实现的字段集合：': (
        '能源法字段分为「固定燃料」「其他燃料（动态行/手动新增）」「净购入电量/热力」「过程排放（动态行/手动新增）」等部分。'
        '能源品种与过程排放项依据《操作指引》附2 Excel 全量配置（按行业差异化展示）；具体字段由对应行业的方法模板决定。'
        '填报页标 <span class="req">*</span> 的为提交必填，其余字段有什么填什么，非必填。除报告法外，能源法不提供附件上传。以下列出系统内已实现的字段示例：'
    ),
    '产品法字段完全由对应行业+方法模板的 product.fields 配置决定，不同行业字段完全不同。以下列出通用字段结构：': (
        '产品法字段依据《操作指引》附2 Excel 按行业配置主要产品与细分项，并支持「其他产品（手动新增）」动态行；'
        '不同行业字段由对应行业+方法模板的 product.fields 决定。填报字段非必填（有什么填什么），除报告法外不提供附件上传。以下列出通用字段结构：'
    ),
    '1、任务名称、核算年度、数据行业范围、行业范围、组织范围、余额口径、数据收集截止日期为必填。': (
        '1、任务名称、核算年度、数据行业范围、行业范围、组织范围、余额口径、数据收集截止日期、分行审批截止日期为必填。'
    ),
    '2、行业范围=自定义时须至少选择一个四级行业代码，否则提示「请至少选择一项行业」。': (
        '2、行业范围=自定义时须至少选择一个四级行业代码，否则提示「请至少选择一项行业」。'
        '\n3、数据收集截止日期须早于分行审批截止日期，否则提示「数据采集截止日期须早于分行审批截止日期」。'
    ),
    '3、组织范围须至少选择全行或一个一级分行。': '4、组织范围须至少选择全行或一个一级分行。',
    '7、填写余额口径、数据收集截止日期、可选的分行审批截止日期。': (
        '7、填写余额口径、数据收集截止日期、分行审批截止日期（须晚于数据收集截止日期）。'
    ),
    '1、校验规则同「新建核算任务」。': '1、校验规则同「新建核算任务」（含数据采集截止日期须早于分行审批截止日期）。',
    '功能描述：审核人查看填报内容，可通过/退回；分行通过时须选定最终采用方法；可调整归属行业与因子。': (
        '功能描述：审核人查看填报内容，可通过/退回；分行通过时须选定最终采用方法（各方法选项展示主体排放预览或「缺少数据」）；'
        '可在审核调整区修改归属行业、选用因子及因子数值。'
    ),
    '4、【退回至客户经理】须填退回原因；任务状态更新为「已退回」，客户经理可重新填报。': (
        '4、【退回至客户经理】须选择退回对象并填写退回原因；任务状态更新为「已退回」，客户经理可重新填报。'
    ),
    '5、【本级修正】：分行绿金岗可直接改数再提交，无需退回。': (
        '5、【修改】：滚动定位至「审核调整」区域，可修改归属行业与排放因子；'
        '【保存】：保存审核调整（归属行业、选用因子、因子数值），不结束审核流程。'
    ),
    '界面设计：审核详情：提示条+调整面板+填报内容只读页签+【审核通过】【退回至客户经理】【本级修正】【取消】。': (
        '界面设计：审核详情：提示条+审核调整面板+填报内容/审批流程页签+'
        '【审核通过】【退回至客户经理】【保存】【修改】【取消】。'
    ),
    '2、可调整归属行业、适用因子并重算排放。': (
        '2、审核调整区可修改归属行业；「选用因子」指定因子库条目（默认预选人行口径内置因子，'
        '下拉分「人行口径（系统默认）」与「我行自定义（可选覆盖）」两组）；'
        '「因子数值」为核算采用的因子值，默认取自所选因子，审核人员可手工覆写；切换选用因子时因子数值自动同步。'
    ),
    '3、多种方法并存时，【审核通过】前须弹窗选择最终采用方法。': (
        '3、多种方法并存时，【审核通过】前须弹窗选择最终采用方法；弹窗各方法选项括号内展示主体排放预览或「缺少数据」。'
    ),
    '因子分内置（只读）与自定义（可维护）；以口径标签（人行口径/行自定义口径）区分，不设因子版本年度。': (
        '因子分内置（人行/指引内置，可编辑/删除，操作前二次确认）与自定义（可维护）；'
        '以口径标签（人行口径/行自定义口径）区分，不设因子版本年度。'
        '内置因子来源列展示人行附2表号（如「人行2-1C」），不再展示 Excel 工作表连写表号（如 2-1CC）。'
    ),
    '1、内置因子不可编辑删除。': (
        '1、内置因子支持【编辑】【删除】（删除前二次确认，提示不可恢复）；自定义因子同。'
    ),
    '3、内置因子：【复制为自定义】【查看】；自定义：【编辑】【删除】。': (
        '3、内置因子与自定义因子均支持【编辑】【删除】【复制】【查看】。'
    ),
    '功能描述：维护自定义因子；选择方法后联动专属字段。': (
        '功能描述：维护内置或自定义因子；选择方法后联动专属字段。'
    ),
    '1、【新增因子】或编辑自定义因子进入表单。': '1、【新增因子】或编辑因子进入表单。',
    '2、方法联动字段：能源法→排放源类型等；产品法→主要产品+口径；经济法→国标行业。': (
        '2、行业大类下拉分两组：「人行八大高碳」「我行主要行业」（GB/T 4754 大类名称），另含「其他」；'
        '方法联动字段：能源法→排放源类型等；产品法→主要产品+口径；经济法→GB/T 4754 四级行业。'
    ),
    '3、【保存】后标记为自定义因子。【复制为自定义】从内置复制。': (
        '3、新增保存后标记为自定义因子；编辑内置因子保存后仍为人行/指引内置因子。'
    ),
    '文案统一「数据收集」「退回」，不用「补录」「驳回」。': (
        '文案统一「数据收集」「退回」，不用「补录」「驳回」。'
        '审批流程页签展示：总行派发→客户经理填报→分行初审→总行终审（总行发起任务时）；'
        '「提交审核」环节文案为「客户经理填报」。'
    ),
    '1、退回原因必填。': '1、退回时须选择退回对象并填写退回原因。',
    '1、来源说明必填；内置因子仅可复制。': '1、来源说明必填。',
    '2、政府/权威报告法：核查=是时须上传附件（每个页签最多3个，单文件≤20MB，pdf/doc/xls/png 等）。': (
        '2、政府/权威报告法：核查=是时须上传附件（每个页签最多3个，单文件≤20MB，pdf/doc/xls/png 等）。'
        '\n3、除报告法外，能源法、产品法、经济法等方法不提供附件上传。'
    ),
    '3、经济法-营收法只读。': '4、经济法-营收法只读。',
}


def replace_paragraph_text(doc: Document) -> int:
    changed = 0
    for para in doc.paragraphs:
        text = para.text
        if text in PARA_REPLACEMENTS:
            para.text = PARA_REPLACEMENTS[text]
            changed += 1
    return changed


def replace_in_tables(doc: Document, old: str, new: str) -> int:
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)
                    changed += 1
    return changed


def delete_table_row(table, row_idx: int) -> None:
    row = table.rows[row_idx]
    tbl = table._tbl
    tbl.remove(row._tr)


def update_tables(doc: Document) -> None:
    # 修订记录
    rev = doc.tables[0]
    rev.add_row()
    cells = rev.rows[-1].cells
    cells[0].text = 'v1.0'
    cells[1].text = '2026.06.30'
    cells[2].text = REVISION_DESC
    cells[3].text = '方丹阳、闫娜'
    cells[4].text = ''

    # 新建任务：分行审批截止日期必填
    t16 = doc.tables[16]
    for row in t16.rows:
        if row.cells[0].text.strip() == '分行审批截止日期':
            row.cells[2].text = '是'
            row.cells[6].text = '须晚于数据收集截止日期；到期前 T-5、T-3 各发送一次企微提醒。'

    # 能源法：取消附件行；部分字段改非必输
    t10 = doc.tables[10]
    for row in t10.rows:
        name = row.cells[0].text.strip()
        if name in ('所属电网', '净购入电量（MWh）'):
            row.cells[2].text = '否'
        if name.startswith('其他燃料'):
            row.cells[6].text = (row.cells[6].text + '；支持手动新增').strip('；')
        if name.startswith('过程排放'):
            row.cells[6].text = (row.cells[6].text + '；按附2全行业配置').strip('；')
    for i, row in enumerate(t10.rows):
        if row.cells[0].text.strip() == '能源法附件':
            delete_table_row(t10, i)
            break

    # 产品法：取消附件；产量字段非必输
    t11 = doc.tables[11]
    for row in t11.rows:
        name = row.cells[0].text.strip()
        if '产量' in name:
            row.cells[2].text = '否'
            row.cells[6].text = (row.cells[6].text + '；有什么填什么').strip('；')
        if name == '主要产品名称/类型':
            row.cells[6].text = (row.cells[6].text + '；支持手动新增其他产品').strip('；')
    for i, row in enumerate(t11.rows):
        if row.cells[0].text.strip() == '产品法附件':
            delete_table_row(t11, i)
            break

    # 审核详情输入项
    t38 = doc.tables[38]
    for row in t38.rows:
        if row.cells[0].text.strip() == '退回原因':
            row.cells[0].text = '退回对象'
            row.cells[1].text = '单选'
            row.cells[2].text = '条件必填'
            row.cells[6].text = '退回时须选择客户经理/提交人等对象'
        if row.cells[0].text.strip() == '适用排放因子':
            row.cells[0].text = '选用因子'
            row.cells[6].text = '默认匹配人行口径内置因子；可选我行自定义因子覆盖'
    extra = t38.add_row().cells
    extra[0].text = '因子数值'
    extra[1].text = '数字输入'
    extra[2].text = '否'
    extra[6].text = '默认取自选用因子，审核人员可覆写'
    reason_row = t38.add_row().cells
    reason_row[0].text = '退回原因'
    reason_row[1].text = '文本域'
    reason_row[2].text = '条件必填'
    reason_row[3].text = '500'
    reason_row[6].text = '退回时必填'

    # 新增因子表单
    t46 = doc.tables[46]
    for row in t46.rows:
        if row.cells[0].text.strip() == '行业（GB四级）':
            row.cells[0].text = '行业大类'
            row.cells[1].text = '下拉框'
            row.cells[5].text = '人行八大高碳/我行主要行业/其他'
            row.cells[6].text = '分组展示：人行八大高碳、我行主要行业（GB/T 4754 大类）；经济法另含 GB 四级行业搜索'


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f'Source not found: {SRC}')
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    n_para = replace_paragraph_text(doc)
    update_tables(doc)
    n_tbl = replace_in_tables(doc, '复制为自定义', '复制')
    doc.save(str(DST))
    print(f'Wrote {DST}')
    print(f'  paragraph replacements: {n_para}')
    print(f'  table text replacements: {n_tbl}')


if __name__ == '__main__':
    main()
