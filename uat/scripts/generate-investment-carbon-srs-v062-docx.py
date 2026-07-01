#!/usr/bin/env python3
"""生成投融资碳核算模块需求规格说明书 v0.62
v0.62：参照《需规修改样例-新建核算任务.docx》统一第4章功能点颗粒度与表述风格；
对齐本地 HTML 原型（http://127.0.0.1:8765/）实际页面字段、按钮文案与交互流程。"""
from pathlib import Path
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from srs_v062_section4 import build_section4
from srs_v062_screenshots import FEATURE_UI_SCREENSHOTS, SHOT_DIR, shot_path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_COLOR_INDEX

ROOT     = Path("/Users/fangdanyang/Desktop/HUAXIA BANK")
TEMPLATE = ROOT / "需规模版.docx"
OUT      = ROOT / "docs" / "华夏银行绿金系统-投融资碳核算模块_需求规格说明书-v0.62.docx"

FONT_BODY  = "宋体"
FONT_TITLE = "微软雅黑"
HEADING_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(12)}


# ─── helpers ──────────────────────────────────────────────────────────────────

def set_run_font(run, font_name, size=None, bold=False, color=None, highlight=None):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def strip_paragraph_numbering(paragraph):
    ppr = paragraph._element.pPr
    if ppr is None:
        return
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def clear_heading_style_numbering(doc):
    """移除模板 Heading 样式自带的多级列表编号，避免与正文手工编号重复。"""
    for i in range(1, 10):
        try:
            style = doc.styles[f"Heading {i}"]
        except KeyError:
            continue
        ppr = style._element.pPr
        if ppr is None:
            continue
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            ppr.remove(num_pr)
        style.font.name = FONT_TITLE
        style.font.bold = True
        rpr = style._element.rPr
        if rpr is not None:
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:ascii"), FONT_TITLE)
            rfonts.set(qn("w:hAnsi"), FONT_TITLE)
            rfonts.set(qn("w:eastAsia"), FONT_TITLE)


def setup_document_fonts(doc):
    clear_heading_style_numbering(doc)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    rpr = normal._element.rPr
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), FONT_BODY)
        rfonts.set(qn("w:hAnsi"), FONT_BODY)
        rfonts.set(qn("w:eastAsia"), FONT_BODY)
        normal.font.size = Pt(10.5)


def _run(p, text, bold=False, size=None, color=None, highlight=None):
    r = p.add_run(text)
    return set_run_font(r, FONT_BODY, size or Pt(10.5), bold=bold, color=color, highlight=highlight)


def hp(doc, text, level):
    p = doc.add_heading(text, level=level)
    strip_paragraph_numbering(p)
    for r in p.runs:
        set_run_font(r, FONT_TITLE, HEADING_SIZES.get(level, Pt(12)), bold=True)
    return p


def hp_feat(doc, text):
    """功能点标题：不加自动编号，宋体加粗。"""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(text), FONT_BODY, Pt(12), bold=True)
    return p

def np(doc, text="", style="Normal"):
    p = doc.add_paragraph(style=style)
    if text: _run(p, text, size=10.5)
    return p

def fp(doc, text):
    p = doc.add_paragraph(style="Normal Indent")
    if text: _run(p, text, size=10.5)
    return p

def lp(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    if text: _run(p, text, size=10.5)
    return p

def pending(doc, text):
    p = doc.add_paragraph(style="Normal")
    _run(p, "【待确认】" + text, highlight=WD_COLOR_INDEX.YELLOW)
    return p

def note(doc, text):
    p = doc.add_paragraph(style="Normal")
    _run(p, "【说明】" + text, color=RGBColor(0x26, 0x74, 0xc1))
    return p

def set_cell_bg(cell, fill="D9EAF7"):
    tcp = cell._tc.get_or_add_tcPr()
    shd = tcp.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcp.append(shd)
    shd.set(qn("w:fill"), fill)

def add_field_table(doc, rows, title_cols=None):
    if title_cols is None:
        title_cols = ["字段名", "类型", "是否必输", "长度", "默认值", "输入限制（或数据字典）", "说明"]
    t = doc.add_table(rows=1, cols=len(title_cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(title_cols):
        hdr[i].text = h
        set_cell_bg(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, FONT_BODY, Pt(9), bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, FONT_BODY, Pt(9))
    return t

def add_plain_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, FONT_BODY, Pt(10), bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, FONT_BODY, Pt(10))
    return t

def add_code_block(doc, lines):
    """伪代码/流程文本块"""
    for line in lines:
        p = doc.add_paragraph(style="Normal Indent")
        r = p.add_run(line)
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        r.font.size = Pt(9)

def _numbered(items):
    return [f"{i}、{t.lstrip('0123456789、. ')}" for i, t in enumerate(items, 1)]


def add_ui_screenshots(doc, shots):
    """在界面设计段落下方插入完整页面截图。"""
    if not shots:
        return
    for name, caption in shots:
        path = shot_path(name)
        if path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Cm(16))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                set_run_font(r, FONT_BODY, Pt(9), color=RGBColor(0x60, 0x60, 0x60))
        else:
            np(doc, f"[缺少截图: {name}.png，请运行 npm run manual:screenshots 生成]", style="Normal Indent")


def feat(doc, no, name, desc, input_rows, flow_lines, check_lines,
         output_text, ui_text, rule_lines, data_text=None, output_rows=None,
         step_suffix=None, ui_screenshots=None):
    title = f"{name}-{step_suffix}" if step_suffix else name
    hp_feat(doc, title)
    fp(doc, f"功能编号：{no}")
    fp(doc, f"功能名称：{name}")
    fp(doc, f"功能描述：{desc}")
    fp(doc, "输入项：")
    if input_rows:
        add_field_table(doc, input_rows)
    if check_lines:
        fp(doc, "校验：")
        for c in _numbered(check_lines):
            lp(doc, c)
    fp(doc, "处理流程：")
    for fl in _numbered(flow_lines):
        lp(doc, fl)
    fp(doc, f"输出项：{output_text}")
    if output_rows:
        add_field_table(doc, output_rows, title_cols=["字段名", "类型", "数据种类", "长度/精度"])
    fp(doc, f"界面设计：{ui_text}")
    shots = ui_screenshots if ui_screenshots is not None else FEATURE_UI_SCREENSHOTS.get(no, [])
    add_ui_screenshots(doc, shots)
    if rule_lines:
        fp(doc, "业务规则：")
        for r in rule_lines:
            lp(doc, r)


# ─── 主函数 ───────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document(str(TEMPLATE))
    setup_document_fonts(doc)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Plain Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CR21093-A_华夏银行绿色金融信息管理系统需求_投融资碳核算模块")
    set_run_font(r, FONT_TITLE, Pt(18), bold=True)

    p = doc.add_paragraph(style="Plain Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("需求规格说明书")
    set_run_font(r, FONT_TITLE, Pt(22), bold=True)

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    for line in ["起草人：Cursor", "起草日期：2026.06.29", "文档版本号：v0.62"]:
        p2 = cell.add_paragraph(style="Table_Small")
        _run(p2, line, size=11)
    for p2 in cell.paragraphs:
        if not p2.text.strip():
            p2._element.getparent().remove(p2._element); break

    np(doc, "本文档中所包含的信息属于内部资料，如无华夏银行的书面许可，任何人都无权复制或利用。", style="Normal Indent")

    # ── 修订记录 ──────────────────────────────────────────────────────────────
    hp(doc, "修订记录", 1)
    add_plain_table(doc, ["版本", "日期", "修订描述", "起草/修订人", "复审人"], [
        ["v0.1", "2026.06.10", "依据当前原型生成初稿", "Cursor", ""],
        ["v0.2", "2026.06.10", "参照《需规模版.docx》调整格式，补充输入项/输出项字段表", "Cursor", ""],
        ["v0.3", "2026.06.23", "依据20260617会议纪要全面修订：字段口径更新、业务种类三档终态、截止日期拆分、审核退回逻辑等", "Cursor", ""],
        ["v0.4", "2026.06.24", "对齐最新原型实现：行业范围改为单选组+四级级联、报告法拆分Tab、因子库按组聚合", "Cursor", ""],
        ["v0.5", "2026.06.26", "新增台账管理、基础配置、权限管理；因子库去版本化；数据采集收集状态5档；审核支持调整行业与因子", "Cursor", ""],
        ["v0.6", "2026.06.26",
         "全面补充数据流转逻辑（含端到端流程图、月均余额计算细节、候选→正式字段映射、客户归集合并规则、各核算方法字段全集、归因排放完整公式）；新增基础配置与数据采集联动说明",
         "Cursor", ""],
        ["v0.62", "2026.06.29",
         "参照样例统一第4章颗粒度；对齐原型并嵌入截图；术语中文化；修正目录双重编号并统一标题字体为微软雅黑/宋体",
         "Cursor", ""],
    ])

    hp(doc, "需求对应记录", 1)
    add_plain_table(doc, ["对应需求编号", "版本", "起草/修订人", "说明"], [
        ["CR21093-A", "v0.62", "Cursor",
         "投融资碳核算模块需求规格说明书，v0.62对齐本地原型并统一功能点表述风格"],
    ])
    add_plain_table(doc, ["需求类型", ""], [
        ["需求类型", "□监管部门需求  □产品服务类  ☑管理决策类  □变更维护类  □其他"],
        ["计划情况", "☑本年度需求计划内  □本年度需求计划外"],
    ])

    hp(doc, "目录", 1)
    np(doc, "（请在 Word 中点击「引用 → 目录 → 更新目录」刷新页码；目录编号与正文一致，采用文档内手工章节编号。）")

    # ──────────────────────────────────────────────────────────────────────────
    # 1 引言
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "1引言", 1)
    hp(doc, "1.1目的", 2)
    np(doc, "本需求规格说明书针对华夏银行绿色金融信息管理系统中的投融资碳核算模块，明确系统应具有的功能、数据处理要求、数据流转规则及计算公式，作为业务人员、开发人员、测试人员和文档编写人员的共同依据。", style="Normal Indent")
    np(doc, "v0.62在第4章按《需规修改样例-新建核算任务.docx》统一功能点颗粒度，并对齐本地 HTML 原型实际页面；第3章数据流转与计算公式仍为权威说明。v0.6相比v0.5的主要增补：完整端到端数据流转图；月均余额计算细节；候选→正式字段映射；各核算方法采集字段清单；归因排放公式；方法模板联动机制等。")

    hp(doc, "1.2背景", 2)
    np(doc, "华夏银行需对投融资业务开展碳排放核算，满足人民银行监管报送和内部绿色金融管理分析需要。投融资碳核算能力以功能模块形式嵌入现有绿金系统，不单独建设独立系统。", style="Normal Indent")
    np(doc, "当前原型已覆盖核算任务发起、接口台账同步、候选清单识别、正式清单确认、核算对象边界确认、数据采集、客户经理在线填报、分级审核、排放计算、结果确认、报告生成、台账管理、排放因子库、企业碳账户及基础配置等完整流程环节。")

    hp(doc, "1.3定义", 2)
    add_plain_table(doc, ["术语", "定义"], [
        ["核算任务",       "围绕某一核算年度或专项范围发起的一次投融资碳核算流程。"],
        ["候选清单",       "系统从信贷台账按规则筛选出的拟纳入核算业务集合，以贷款笔为粒度（1笔=1行）。"],
        ["正式清单",       "管理端确认并锁定后的待核算业务清单，继承候选清单字段，状态推进至confirmed。"],
        ["业务种类",       "三档终态：非项目贷款、项目贷款（以项目方式计算）、项目贷款（以非项目方式计算）。中间状态：项目（计算方法待定）。"],
        ["收集任务",       "针对正式清单派发给客户经理的在线填报任务（supplement），用于补充排放数据、项目信息和佐证附件。"],
        ["归集单元",       "数据采集合并层：同一客户（信用代码）下的同类（非项目/项目）贷款合并为一个收集单元，对应一个supplement。"],
        ["主体排放",       "融资主体或项目在核算周期内的碳排放量（tCO₂e），来源为报告法、能源法、产品法、经济法或兜底法。"],
        ["归因排放",       "按本行投融资占比归因至本行业务的碳排放量（tCO₂e），计算公式见3.3节。"],
        ["月均贷款余额",   "核算年度内该笔贷款存续月份的月末余额均值；内部存储单位为万元，展示单位为元（×10000）。"],
        ["华夏存续月份",   "该笔贷款在核算年度内于华夏银行存续的月份数（1~12），由发放日期与核算年度计算得出。"],
        ["月末余额合计",   "月均贷款余额×华夏存续月份，用于归因分母加总。"],
        ["平均资产总额",   "上年末合并报表资产总额与当年末合并报表资产总额之和除以2，单位元，来源信贷接口。"],
        ["DQR",           "数据质量评级：DQR=Σ(单笔排放×质量得分)/Σ单笔排放，按A/B+/B/B-/C五档。"],
        ["企业碳账户",     "以统一社会信用代码为主键的客户维度排放账户；确认结果后自动归集。"],
        ["信贷数据兜底法", "以月均贷款余额×行业排放因子估算主体排放，对应「其他计算法」，DQR质量等级最低。"],
        ["方法模板",       "按行业+业务类型+核算方法维护的采集字段模板，定义收集填报页展示哪些输入字段及计算公式。"],
        ["参数字段库",     "可复用的字段定义集合，供方法模板引用；包含字段名、数据类型、单位、枚举选项等。"],
    ])

    hp(doc, "1.4参考资料", 2)
    for ref in [
        "《商业银行投融资业务碳核算与报告指南》（人行）",
        "《华夏银行绿金系统-投融资碳核算与企业碳账户-项目需求说明书》v0.1",
        "《华夏银行投融资碳核算及企业碳账户需求评审会议纪要》2026年6月17日、2026年6月23日",
        "《数据采集-客户归集合并改造计划》（内部方案稿，2026.06.23）",
        "当前投融资碳核算原型（assets/js/）及自测脚本（test-flow.js、test-full.js）",
    ]:
        np(doc, ref)

    # ──────────────────────────────────────────────────────────────────────────
    # 2 任务概述
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "2任务概述", 1)
    hp(doc, "2.1目标", 2)
    np(doc, "建设嵌入绿金系统的投融资碳核算模块，支持总行或分行完成从核算任务创建、信贷台账同步、清单确认、数据采集、排放计算、结果确认到报告输出和碳账户归集的完整闭环。")
    for item in [
        "支持按核算年度、行业范围和组织范围发起核算任务；行业范围支持参数化配置。",
        "支持从接口管理同步信贷台账，生成候选清单和正式清单；行业识别以贷款投向所属行业为主要筛选依据。",
        "支持按信贷品种判断项目类/非项目类，结合项目信息形成三档业务种类终态。",
        "数据采集环节支持客户归集合并：同一客户（统一社会信用代码）下的同类贷款合并为一个收集单元下发；填报完成后主体排放拆回逐笔归因。",
        "支持系统自动获取格澜等外部接口数据（报告法来源）和行内系统数据（营收、余额等）形成系统底稿。",
        "支持向客户经理发放数据采集任务；收集填报页字段根据方法模板动态展示，不同行业+方法组合呈现不同字段。",
        "支持分行初审时明确最终采用的核算方法，支持退回至客户经理；审核时可调整归属行业及适用因子。",
        "支持排放计算（含DQR）、结果确认、报告导出和企业碳账户归集。",
        "支持台账管理、基础配置（参数字段库/方法模板/行业配置）和权限管理。",
    ]:
        lp(doc, item)

    hp(doc, "2.2约束", 2)
    for item in [
        "项目类信贷品种完整映射关系尚未由业务侧提供；当前按已知品种判断。",
        "数据采集客户归集合并规则（集团主办行、项目主办分行等）需业务侧确认字段来源，详见4.3节待确认清单。",
        "正式生产环境的接口报文、审批状态码、权限模型、报告模板以华夏银行最终确认材料为准。",
        "审批流、待办中心、组织权限和用户体系不在投融资碳核算模块内重复建设，沿用绿金系统既有能力。",
    ]:
        pending(doc, item)

    # ──────────────────────────────────────────────────────────────────────────
    # 3 端到端数据流转总览（v0.6 新增）
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "3端到端数据流转与核心算法", 1)

    hp(doc, "3.1完整数据流转流程图", 2)
    np(doc, "下图描述投融资碳核算模块从信贷接口到企业碳账户的完整数据流转路径：")
    add_code_block(doc, [
        "┌─────────────────────────────────────────────────────────────────────────────┐",
        "│                    投融资碳核算 — 端到端数据流转                               │",
        "└─────────────────────────────────────────────────────────────────────────────┘",
        "",
        " [信贷核心系统]                                                                 ",
        "  逐月按批次推送 ──→ [接口管理] 按年度汇总成功批次                               ",
        "                                  │                                            ",
        "                         ▼ 点击「从接口同步台账」                                ",
        " [候选清单] ←── 按默认规则筛选（余额≥500万、排除小微/个人/境外/非高碳）           ",
        "    │  每行 = 1笔贷款                                                           ",
        "    │  字段：客户名称、信用代码、贷款账号、信贷品种、一级分行、经办行、            ",
        "    │         月均贷款余额（万元）、月末余额合计（万元）、华夏存续月份、           ",
        "    │         年报营业收入（万元）、合并报表资产总额（万元）、上年末资产（万元）、 ",
        "    │         贷款投向所属行业代码/大类、企业所属行业代码、主办客户经理            ",
        "    │                                                                           ",
        "    ▼ 勾选纳入 → 点击「生成正式清单」                                            ",
        " [正式清单] ← 继承候选清单全部字段，status=draft → confirmed（锁定后）          ",
        "    │  同时：建档企业碳账户空户（status=active）                                 ",
        "    │                                                                           ",
        "    ▼ 点击「调取接口数据」                                                       ",
        "    ├─→ [格澜接口] 按信用代码查询报告法主体排放 → 写入 gelanEntityEmission       ",
        "    │              预填至客户经理填报页「其他来源报告法」Tab（可编辑）            ",
        "    ├─→ [经济法直算] 非项目类 → 主体排放 = 年报营业收入(万元) × 行业因子          ",
        "    │                           归因排放 = 主体排放 × 月均余额 / 平均资产总额    ",
        "    │                                                                           ",
        "    ▼ 「发放收集任务」                                                           ",
        "    ├─→ [客户归集合并] 同一信用代码下的非项目笔合并→1个非项目收集单元             ",
        "    │                  同一信用代码下的项目笔合并→1个项目收集单元                 ",
        "    │                  非项目跨多分行→下发集团主办行；单分行→下发该分行           ",
        "    │                  项目类→下发项目主办分行                                   ",
        "    │                                                                           ",
        "    ▼ [收集任务(supplement)] 1个归集单元=1个supplement                          ",
        "    ├─→ [客户经理填报] 方法模板动态字段                                          ",
        "    │     ├ 报告法-政府/权威来源：scope1/2排放、数据来源、核查情况、附件          ",
        "    │     ├ 报告法-其他来源（格澜预填可编辑）：排放总量、scope1/2、来源           ",
        "    │     ├ 能源法：固定燃料、其他燃料（动态行）、净购入电量、过程排放             ",
        "    │     ├ 产品法：主要产品产量（按模板字段）                                    ",
        "    │     ├ 经济活动法-营收法：只读，系统预填（不可编辑）                         ",
        "    │     └ 项目信息（项目类）：是否可提供→确定业务种类终态                       ",
        "    │                                                                           ",
        "    ▼ 分行审核（初审/终审）                                                      ",
        "    ├─→ 审核通过：选定最终采用方法 → 可调整归属行业与排放因子                     ",
        "    │                                                                           ",
        "    ▼ [排放计算] 主体排放已确定 → 逐笔归因拆分                                   ",
        "    │  非项目：E_归因 = E_主体 × (该笔月均余额) / (归集单元合计平均资产总额)     ",
        "    │  项目（以项目方式）：E_归因 = E_项目 × (该笔月均余额) / 项目总投资         ",
        "    │  兜底法：E_主体 = 月均余额(万) × 行业因子；E_归因 = E_主体（全额归因）     ",
        "    │  DQR 计算：各笔排放加权质量得分→年度综合评级                               ",
        "    │                                                                           ",
        "    ▼ 【确认结果】                                                               ",
        "    └─→ [企业碳账户] 写入排放记录（逐笔）→ 归集至对应账户 annualProfiles         ",
        "                      账户主键 = 统一社会信用代码（同一主体一个账户）             ",
    ])

    hp(doc, "3.2月均贷款余额计算细节", 2)
    np(doc, "候选清单同步时，系统从信贷接口逐笔取月末余额原始数据，按以下规则计算候选清单和正式清单中的余额字段：")
    hp(doc, "3.2.1华夏存续月份", 3)
    np(doc, "华夏存续月份（huaxiaTenureMonths）= 该笔贷款在核算年度内于华夏银行存续的自然月数，取值范围 1~12。", style="Normal Indent")
    add_plain_table(doc, ["场景", "规则", "示例"], [
        ["发放日期早于核算年度", "存续月份 = 12（全年存续）", "2023年发放，核算年度=2024 → 12"],
        ["发放日期在核算年度内", "存续月份 = 12 - 发放月份 + 1", "2024-03月发放 → 12-3+1=10"],
        ["发放日期晚于核算年度", "存续月份 = 1（最小值保护）", "2025年发放，核算年度=2024 → 1"],
        ["发放日期为空", "默认 = 12", "接口未提供发放日期 → 12"],
    ])
    add_code_block(doc, [
        "// 伪代码",
        "function computeHuaxiaTenureMonths(disbursementDate, accountingYear) {",
        "  if (disbursementDate 为空) return 12;",
        "  const [dy, dm] = parseYearMonth(disbursementDate);",
        "  if (dy < accountingYear)  return 12;",
        "  if (dy > accountingYear)  return 1;          // 当年发放但已超年末",
        "  return max(1, 12 - dm + 1);                  // 当年发放",
        "}",
    ])

    hp(doc, "3.2.2月均贷款余额（月均余额）", 3)
    np(doc, "月均贷款余额（avgMonthlyBalance）= 月末余额合计 / 华夏存续月份", style="Normal Indent")
    np(doc, "月末余额合计（monthEndBalanceSum）= 各存续月份月末余额之和。接口未提供逐月月末数据时，以接口提供的avgMonthlyBalance作为初始值：", style="Normal Indent")
    add_code_block(doc, [
        "// 接口提供逐月数据时：",
        "monthEndBalanceSum = sum(月末余额Jan..月末余额Dec [有余额的月份]);",
        "avgMonthlyBalance  = monthEndBalanceSum / huaxiaTenureMonths;",
        "",
        "// 接口仅提供均值时（当前演示场景）：",
        "monthEndBalanceSum = avgMonthlyBalance × huaxiaTenureMonths;",
        "// avgMonthlyBalance 直接用接口值，不再重算",
        "",
        "// 内部存储单位：万元（小数）",
        "// 展示单位：元 → formatLedgerAmountYuan(wan) = wan × 10000 格式化",
    ])
    note(doc, "候选清单中展示「月末余额（元）」= monthEndBalanceSum × 10000；正式清单中展示「月均贷款余额（元）」= avgMonthlyBalance × 10000。两者使用不同字段，切勿混淆。")

    hp(doc, "3.2.3平均资产总额", 3)
    np(doc, "平均资产总额（avgTotalAssets）=（上年末合并报表资产总额 + 当年末合并报表资产总额）/ 2，单位元。", style="Normal Indent")
    add_code_block(doc, [
        "avgTotalAssets = (prevYearTotalAssets + currentYearTotalAssets) / 2;",
        "// prevYearTotalAssets：来源信贷接口，字段名 prevYearTotalAssets（万元）",
        "// currentYearTotalAssets：来源信贷接口，字段名 totalAssets（万元）",
        "// 候选清单展示「合并报表资产总额（元）」= currentYearTotalAssets × 10000",
        "// 正式清单展示「平均资产总额（元）」   = avgTotalAssets × 10000",
    ])

    hp(doc, "3.3候选清单→正式清单字段映射", 2)
    np(doc, "点击「生成正式清单」时，系统按1:1从候选清单创建正式清单记录，字段继承规则如下：")
    add_plain_table(doc, ["字段", "候选清单展示名", "正式清单展示名", "存储字段", "单位/换算", "备注"], [
        ["客户名称",           "客户名称",               "客户名称",               "customerName",          "文本",      ""],
        ["统一社会信用代码",   "统一社会信用代码",         "统一社会信用代码",         "creditCode",            "文本",      "企业碳账户主键"],
        ["贷款账号",           "贷款账号",               "贷款账号",               "loanAccount",           "文本",      ""],
        ["信贷品种",           "信贷品种",               "信贷品种",               "loanType",              "文本",      ""],
        ["业务种类",           "业务种类",               "业务种类",               "accountingType",        "枚举",      "由信贷品种+项目信息推导"],
        ["一级分行",           "一级分行",               "一级分行",               "tier1Branch",           "文本",      ""],
        ["经办行",             "经办行",                 "经办行",                 "handlingBranch",        "文本",      ""],
        ["主办客户经理",       "主办客户经理",             "主办客户经理",             "manager",               "文本",      ""],
        ["华夏存续月份",       "（内部字段）",             "（内部字段）",             "huaxiaTenureMonths",    "整数/月",   "候选同步时计算"],
        ["月末余额合计",       "月末余额（元）",           "（内部参考）",             "monthEndBalanceSum",    "万元→展示元", "候选清单显示此字段"],
        ["月均贷款余额",       "（内部字段）",             "月均贷款余额（元）",       "avgMonthlyBalance",     "万元→展示元", "正式清单显示此字段"],
        ["年报营业收入",       "年报营业收入（元）",       "年报营业收入（元）",       "operatingRevenue",      "万元→展示元", ""],
        ["上年末合并报表资产", "（内部字段）",             "（内部字段）",             "prevYearTotalAssets",   "万元",      "用于计算avgTotalAssets"],
        ["当年末合并报表资产", "合并报表资产总额（元）",   "（内部参考）",             "totalAssets",           "万元→展示元", "候选清单显示此字段"],
        ["平均资产总额",       "（内部字段）",             "平均资产总额（元）",       "avgTotalAssets",        "万元→展示元", "正式清单显示此字段"],
        ["贷款投向行业代码",   "贷款投向所属行业",         "贷款投向所属行业",         "gbIndustryCode",        "文本/GB四级", "筛选主依据"],
        ["企业所属行业代码",   "企业所属行业",             "企业所属行业",             "industryMajor",         "文本",      "辅助展示字段"],
        ["项目明细",           "（展开查看）",             "（展开查看）",             "projectDetails",        "JSON数组",  "项目类业务"],
    ])

    hp(doc, "3.4数据采集客户归集合并规则", 2)
    np(doc, "正式清单锁定后，「发放收集任务」前系统自动执行客户归集合并（buildCollectGroups），将同一客户的多笔贷款合并为归集单元（collectGroup），以便按客户维度统一下发收集任务、填报和审核。")
    note(doc, "台账逐笔数据不动；合并只发生在「数据采集下发」层。归因排放仍逐笔计算，以保证监管报送精度。")

    hp(doc, "3.4.1归集维度", 3)
    add_plain_table(doc, ["归集维度", "规则"], [
        ["客户主键",     "统一社会信用代码（creditCode），同一主体不同贷款账号合并"],
        ["归集桶",       "非项目（non_project）和项目（project）分开，同一客户最多2个收集单元"],
        ["桶判定",       "非项目贷款→非项目桶；项目类（含待定）→项目桶"],
    ])

    hp(doc, "3.4.2下发分行规则", 3)
    add_plain_table(doc, ["场景", "下发对象", "说明"], [
        ["非项目，仅1个一级分行",    "该分行",       "单分行直接下发"],
        ["非项目，跨多个一级分行",   "集团主办行",    "需接口/CRM提供groupLeadBranch字段；待确认缺失时默认规则"],
        ["项目类",                   "项目贷款主办分行", "需字段projectLeadBranch；多项目跨分行时优先级待确认"],
    ])
    pending(doc, "集团主办行、项目贷款主办分行字段的来源（信贷接口/CRM/人工指定）及缺失时默认值，待业务与技术确认。")

    hp(doc, "3.4.3金额字段归集方式", 3)
    add_plain_table(doc, ["字段", "归集方式", "说明"], [
        ["月均贷款余额",       "求和（Σ各笔）",                 "归集单元合计余额，用于后续逐笔归因分母"],
        ["年报营业收入",       "客户级去重（取首笔或最大值）",   "合并报表口径，不重复累加"],
        ["平均资产总额",       "客户级去重",                     "合并报表口径，不重复累加"],
        ["项目类字段",         "保留projectDetails列表",         "多项目合并至同一收集单元"],
    ])

    hp(doc, "3.4.4归集核算行业规则（非项目）", 3)
    add_code_block(doc, [
        "// 非项目归集单元核算行业确定规则：",
        "const investCodes = unique(nonProject.map(r => r.gbIndustryCode));",
        "if (investCodes.length > 1) {",
        "  // 多个投向行业 → 取客户所属行业（企业所属行业，industryMajor）",
        "  accountingIndustryCode   = pickCustomerIndustry(creditCode, nonProject);",
        "  accountingIndustrySource = 'customer';",
        "} else {",
        "  // 单一投向 → 用贷款投向行业",
        "  accountingIndustryCode   = investCodes[0];",
        "  accountingIndustrySource = 'invest';",
        "}",
        "// 项目类 → 仍按项目明细各自行业填报，不做行业合并",
    ])
    pending(doc, "非项目单一投向时，核算行业使用投向还是客户所属行业，请业务侧确认统一口径。")

    hp(doc, "3.4.5归集单元数据结构", 3)
    add_code_block(doc, [
        "collectGroup = {",
        "  id: 'G{taskId}_{seq}',",
        "  taskId,",
        "  creditCode,",
        "  customerName,",
        "  bucket: 'non_project' | 'project',",
        "  memberFormalIds: ['F001', 'F002', ...],  // 关联逐笔正式清单",
        "  memberCount: 3,",
        "  dispatchBranch: '上海分行',",
        "  dispatchRule: 'single_branch' | 'group_lead' | 'project_lead',",
        "  accountingIndustryCode: 'C3011',",
        "  accountingIndustrySource: 'customer' | 'invest' | 'project',",
        "  aggregatedBalance: 12500,              // 合计月均余额（万元）",
        "  revenue: 68000,                        // 客户级营业收入（万元，去重）",
        "  totalAssets: 320000,                   // 客户级平均资产总额（万元，去重）",
        "  groupLeadBranch: '北京分行',",
        "  projectLeadBranch: '深圳分行',",
        "  status: 'pending' | 'dispatched' | 'completed',",
        "  supplementId: 'S...'                   // 1:1关联supplement",
        "}",
    ])

    hp(doc, "3.5各核算方法采集字段全集", 2)
    np(doc, "客户经理在线收集填报页（supplement）中，根据行业+业务类型+方法模板动态展示字段。以下按方法类型列出所有可能采集字段（具体行业模板可能为子集）：")

    hp(doc, "3.5.1报告法—政府/权威来源（reportAuthority）", 3)
    add_field_table(doc, [
        ("报告法数据来源", "下拉框", "是", "50", "EPD认证报告", "ESG报告/环境信息披露/政府核查/CCER/第三方核查", ""),
        ("该数据是否经政府/第三方核查", "下拉框", "是", "10", "是", "是/否", "是→须上传附件"),
        ("温室气体排放总量（tCO₂e）", "数字输入", "是", "20,4", "", ">0", "scope1+scope2合计"),
        ("范围一排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", "直接排放"),
        ("范围二排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", "间接排放"),
        ("核查报告期间起", "日期", "否", "10", "", "YYYY-MM-DD", ""),
        ("核查报告期间止", "日期", "否", "10", "", "YYYY-MM-DD", ""),
        ("报告附件", "文件上传", "条件必填", "—", "", "pdf/doc/xls/png/jpg，最多3个，单文件≤20MB", "核查=是时必传"),
    ])

    hp(doc, "3.5.2报告法—其他来源（reportOther，含格澜预填）", 3)
    add_field_table(doc, [
        ("报告法数据来源", "下拉框", "是", "50", "格澜数据-各地区企业环境信息披露平台", "ESG报告/格澜数据-.../行业协会/其他", "格澜数据预填此字段"),
        ("该数据是否经政府/第三方核查", "下拉框", "是", "10", "否", "是/否", ""),
        ("温室气体排放总量（tCO₂e）", "数字输入", "是", "20,4", "", ">0", "格澜预填，可编辑"),
        ("范围一排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", ""),
        ("范围二排放（tCO₂e）", "数字输入", "否", "20,4", "", ">0", ""),
        ("报告附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    note(doc, "格澜接口成功返回数据后，系统自动将数据写入「其他来源报告法」Tab，并标注来源为「格澜数据-各地区企业环境信息披露平台」。客户经理可在填报页查看并修改数值或来源说明，但不可删除该预填记录。")

    hp(doc, "3.5.3物理活动法—能源法（energy，字段由方法模板决定）", 3)
    np(doc, "能源法字段分为「固定燃料」「其他燃料（动态行）」「净购入电量/热力」「过程排放（动态行）」四个部分，具体字段由对应行业的方法模板配置决定。以下列出系统内已实现的字段集合：")
    add_field_table(doc, [
        ("固定燃料-原煤消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", "方法模板fuelFixed"),
        ("固定燃料-焦炭消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-燃料油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-天然气消耗量（万立方米）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-汽油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("固定燃料-柴油消耗量（万吨）", "数字输入", "按模板", "20,4", "", ">0", ""),
        ("其他燃料（动态行）-燃料品种", "下拉框", "否", "50", "", "原油/燃料油/汽油/柴油/液化石油气/液化天然气/焦炉煤气/高炉煤气/转炉煤气/炼厂干气/其他煤气/无", "可增删行，最多20行"),
        ("其他燃料（动态行）-消耗量（吨或万立方米）", "数字输入", "否", "20,4", "", ">0", "与品种行对应"),
        ("所属电网", "下拉框", "是", "50", "全国平均", "全国平均/华北/华东/华中/南方/东北/西北", "由模板gridOptions决定"),
        ("净购入电量（MWh）", "数字输入", "是", "20,4", "", "≥0", ""),
        ("净购入热力（GJ）", "数字输入", "否", "20,4", "", "≥0", "仅部分模板展示"),
        ("过程排放-脱硫试剂类型", "下拉框", "否", "50", "", "石灰石/白云石/方解石/其他", "动态行，按模板processBlocks"),
        ("过程排放-脱硫试剂消耗量（吨）", "数字输入", "否", "20,4", "", ">0", ""),
        ("过程排放-碳酸盐类型", "下拉框", "否", "50", "", "石灰石/白云石/其他", "动态行，按模板"),
        ("过程排放-碳酸盐消耗量（吨）", "数字输入", "否", "20,4", "", ">0", ""),
        ("能源法附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    np(doc, "能源法主体排放计算公式（演示原型暂不自动试算，以业务确认版为准）：")
    add_code_block(doc, [
        "E_主体 = Σ(燃料消耗量 × 对应排放因子)         // 燃料燃烧",
        "       + Σ(脱硫/碳酸盐消耗量 × 工艺因子)       // 过程排放",
        "       + 净购入电量(MWh) × 区域电网因子         // 外购电力",
        "       + 净购入热力(GJ) × 热力因子;             // 外购热力（部分行业）",
    ])

    hp(doc, "3.5.4物理活动法—产品法（product，字段由方法模板决定）", 3)
    np(doc, "产品法字段完全由对应行业+方法模板的 product.fields 配置决定，不同行业字段完全不同。以下列出通用字段结构：")
    add_field_table(doc, [
        ("主要产品名称/类型", "文本/模板内置", "按模板", "100", "", "", "模板中固定展示"),
        ("主要产品产量（万吨）", "数字输入", "是", "20,4", "", ">0", "按模板分组展示"),
        ("细分产品产量（万吨）", "数字输入", "按模板", "20,4", "", ">0", "部分行业有细分"),
        ("产品法附件", "文件上传", "否", "—", "", "同上", ""),
    ])
    add_code_block(doc, [
        "E_主体 = Σ(产品产量 × 产品碳排放因子);   // 各产品类型分别计算后求和",
        "// 产品碳排放因子来源：排放因子库 caliberTag=pbo（人行口径）或 caliberTag=bank（行自定义）",
    ])

    hp(doc, "3.5.5经济活动法—营收法（economy，系统预填只读）", 3)
    np(doc, "经济活动法-营收法字段由系统从正式清单数据预填，客户经理填报页只读，不可编辑。")
    add_field_table(doc, [
        ("年报营业收入（万元）", "只读展示", "—", "20,2", "", "来自正式清单operatingRevenue", "系统自动填入"),
        ("行业排放因子（tCO₂e/万元营收）", "只读展示", "—", "20,6", "", "来自因子库经济活动法行业因子", "系统自动匹配"),
        ("系统计算主体排放（tCO₂e）", "只读展示", "—", "20,4", "", "= 营收 × 因子", "系统计算"),
    ])
    add_code_block(doc, [
        "E_主体 = 年报营业收入(万元) × 行业排放因子(tCO₂e/万元营收);",
        "// 行业因子匹配优先级：GB国标四级代码 > 行业大类",
        "// 若营收为0或缺失 → 该方法无法计算，提示「请提供营业收入」",
    ])

    hp(doc, "3.5.6信贷数据兜底法（credit_fallback，不在收集填报页展示）", 3)
    np(doc, "信贷数据兜底法由管理端直接执行，不在客户经理收集填报页展示。触发条件：仍有记录无主体排放时，点击「信贷数据兜底法」按钮。")
    add_code_block(doc, [
        "E_主体 = 月均贷款余额(万元) × 行业排放因子(tCO₂e/万元余额);",
        "E_归因 = E_主体;   // 全额归因，不按资产占比折算",
        "// 核算方法标记为「其他计算法」；DQR质量等级=5（最低）",
        "// 行业因子取 gbIndustryCode 对应经济活动法因子",
    ])

    hp(doc, "3.6归因排放完整计算公式", 2)
    np(doc, "排放计算环节，系统按业务种类和核算路径分档计算归因排放：")
    add_plain_table(doc, ["业务种类", "主体排放来源", "归因排放公式", "归因占比分母", "说明"], [
        ["非项目贷款",
         "报告法/能源法/产品法/经济法/兜底法",
         "E_归因 = E_主体 × (该笔月均余额/12) / (平均资产总额×12) × 12",
         "融资主体平均资产总额",
         "简化写法：E_归因 = E_主体 × 月均余额 / 平均资产总额"],
        ["项目贷款（以项目方式计算）",
         "客户经理填报项目排放数据",
         "E_归因 = E_项目 × 该笔月均余额 / 项目总投资",
         "项目总投资（projectTotalInvestmentWan×10000）",
         "项目维度归因"],
        ["项目贷款（以非项目方式计算）",
         "同非项目贷款",
         "E_归因 = E_主体 × 月均余额 / 平均资产总额",
         "融资主体平均资产总额",
         "与非项目同公式"],
        ["信贷数据兜底法",
         "月均余额×行业因子",
         "E_归因 = E_主体（全额归因）",
         "不折算",
         "DQR最低"],
    ])
    add_code_block(doc, [
        "// 伪代码",
        "for each formal f in confirmedFormals(taskId) {",
        "  const calc = getCalculation(f.id);",
        "  if (calc.source === 'credit_fallback') {",
        "    f.attributedEmission = calc.entityEmission;   // 全额归因",
        "  } else if (f.accountingType === 'project_as_project') {",
        "    const projectInvest = f.projectTotalInvestmentWan * 10000;",
        "    f.attributedEmission = calc.entityEmission * f.avgMonthlyBalance / projectInvest;",
        "  } else {",
        "    // non_project 或 project_as_non_project",
        "    const avgAssets = f.avgTotalAssets * 10000;   // 换算为元",
        "    const balance   = f.avgMonthlyBalance * 10000; // 换算为元",
        "    f.attributedEmission = calc.entityEmission * balance / avgAssets;",
        "  }",
        "}",
    ])

    hp(doc, "3.7方法模板与参数字段库联动机制", 2)
    np(doc, "基础配置模块（参数字段库+方法模板）与数据采集收集填报的联动流程如下：")
    add_code_block(doc, [
        "配置层（基础配置/方法模板）：",
        "  总行维护「参数字段库」→ 定义可复用字段（名称/类型/单位/枚举等）",
        "  总行维护「方法模板」→ 按行业+业务类型+核算方法组合，",
        "                         引用参数字段库中的字段，配置计算公式和因子绑定",
        "  点击「发布」→ 模板status=published，对应行业+方法组合生效",
        "",
        "展示层（数据采集/收集填报）：",
        "  系统根据 formal.industryMajor + formal.accountingType",
        "  查找 status=published 的方法模板",
        "  若找到 → 按模板字段渲染收集填报页各方法Tab",
        "  若未找到 → 按默认字段集（defaultMode）渲染",
        "",
        "字段渲染规则（以能源法为例）：",
        "  template.methods.energy.fuelFixed  → 固定燃料输入行",
        "  template.methods.energy.otherFuelOptions → 其他燃料下拉选项",
        "  template.methods.energy.gridOptions → 所属电网下拉选项",
        "  template.methods.energy.hasPurchasedHeat → 是否展示净购入热力字段",
        "  template.methods.energy.processBlocks → 过程排放动态行配置",
        "  template.methods.product.fields → 产品法产品字段列表",
    ])
    note(doc, "当前原型已实现玻璃行业、能源行业等方法模板示例，其余行业模板由总行管理员在「方法模板配置」页面维护后发布生效，无需修改系统代码。")

    hp(doc, "3.8主体排放→逐笔归因拆分流程", 2)
    np(doc, "客户归集合并后，supplement填报的主体排放为归集单元（collectGroup）级别，需在排放计算时拆回逐笔归因：")
    add_code_block(doc, [
        "// 归集单元级主体排放已由审核确认",
        "const groupEntityEmission = supplement.fieldData.entityEmission;",
        "",
        "for each formalId in collectGroup.memberFormalIds {",
        "  const f = getFormal(formalId);",
        "  // 非项目：各笔共享同一主体排放，按各笔余额占比归因",
        "  const totalBalance = sum(memberFormals.map(m => m.avgMonthlyBalance));",
        "  f.entityEmission    = groupEntityEmission;",
        "  f.attributedEmission = groupEntityEmission",
        "                        * f.avgMonthlyBalance / f.avgTotalAssets;",
        "  // 注：avgTotalAssets为客户级（去重），各笔相同",
        "}",
    ])

    hp(doc, "3.9DQR数据质量评价", 2)
    add_plain_table(doc, ["核算方法", "质量得分（qualityScore）", "说明"], [
        ["报告法-政府/权威来源（经核查）", "1", "最高质量"],
        ["报告法-政府/权威来源（未核查）", "2", ""],
        ["报告法-其他来源（含格澜）",      "2", ""],
        ["物理活动法-能源法",              "2", ""],
        ["物理活动法-产品法",              "3", ""],
        ["经济活动法-营收法（直算）",      "4", ""],
        ["信贷数据兜底法（其他计算法）",   "5", "最低质量"],
    ])
    add_code_block(doc, [
        "// 年度综合DQR",
        "DQR = Σ(E_归因_i × qualityScore_i) / Σ(E_归因_i);",
        "// DQR等级映射：",
        "// ≤1.5 → A     ≤2.0 → B+    ≤3.0 → B",
        "// ≤4.0 → B-    >4.0 → C",
    ])

    # ──────────────────────────────────────────────────────────────────────────
    # 4 功能需求（与v0.5一致，引用3章中的数据规则）
    # ──────────────────────────────────────────────────────────────────────────
    # ──────────────────────────────────────────────────────────────────────────
    # 4 功能需求（v0.62 对齐样例颗粒度 + 本地原型）
    # ──────────────────────────────────────────────────────────────────────────
    hp(doc, "4需求规定", 1)
    build_section4(doc, hp, np, feat)

    # ── 4.2 外部接口需求 ─────────────────────────────────────────────────────
    hp(doc, "4.2外部接口需求", 2)
    add_plain_table(doc, ["外部接口", "说明", "待确认事项"], [
        ["信贷台账接口", "按月批次推送贷款台账；字段含分行、客户、品种、月均余额、月末余额合计、华夏存续月份、营收、资产（当年末+上年末）、行业、主办客户经理、项目信息（projectDetails）等。", "正式接口字段名/类型；调用频率；错误码；重试；脱敏要求。"],
        ["格澜企业排放接口", "按信用代码查询企业主体碳排放（报告法数据）；返回ghgTotalEmission、scope1/2Emission、reportYear、reportSource等字段。", "生产接口地址；认证方式；请求参数；返回字段；频率与缓存策略。"],
        ["审批与待办接口", "提交审核传入单据类型/业务ID/提交人；接收审批状态回写（通过/退回/作废）。", "正式审批状态码；流程实例ID；退回后解锁字段范围。"],
        ["组织权限接口", "复用绿金系统组织树、用户、角色和数据权限。", "角色枚举和权限控制粒度。"],
        ["附件与文件服务", "收集填报须支持附件上传；病毒扫描、预览、留存年限。", "文件服务接口、大小/数量限制。"],
    ])

    # ── 4.3 待确认事项汇总（v0.6新增完整清单） ──────────────────────────────
    hp(doc, "4.3待确认事项汇总", 2)
    add_plain_table(doc, ["序号", "主题", "待确认内容", "影响范围"], [
        ["C01", "信贷品种→项目类映射",       "项目贷款/一般性固定资产贷款/出口退税账户托管贷款等完整清单及判断逻辑", "业务种类识别"],
        ["C02", "集团主办行字段",             "非项目类跨分行时，集团主办行（groupLeadBranch）字段从信贷接口/CRM/人工指定哪个来源取", "归集合并下发规则"],
        ["C03", "项目贷款主办分行字段",       "projectLeadBranch字段来源及多项目跨分行时的优先级规则（金额最大/首笔/接口指定）", "归集合并下发规则"],
        ["C04", "非项目单投向核算行业",       "非项目归集单元仅1个贷款投向行业时，核算行业用投向行业还是客户所属行业", "归集核算行业"],
        ["C05", "业务种类名称确认",           "三档终态名称是否与人行附件4完全一致（含是否带「贷款」二字）", "业务种类识别/展示"],
        ["C06", "营收/资产客户级字段规则",    "归集单元内多笔贷款，客户级年报营收、合并报表资产总额以哪笔/哪个系统为准（建议合并报表数仓取单次）", "归因排放计算"],
        ["C07", "方法模板发布审批",           "方法模板发布是否需要审批流程，由谁审批", "基础配置/方法模板"],
        ["C08", "格澜接口调取粒度",           "格澜接口是按正式清单笔调取还是按归集单元（creditCode）调取一次", "调取接口数据"],
        ["C09", "报告模板提供",               "监管与管理报告模板、导出维度由行方提供的清单与时间表", "生成报告"],
        ["C10", "碳账户排放结果更正",         "核算结果更正后，碳账户内对应记录是覆盖、冲正还是保留多版本", "企业碳账户"],
        ["C11", "碳账户停用规则",             "停用/注销账户是否仍参与历史统计；注销后是否允许再启用", "企业碳账户状态管理"],
        ["C12", "华夏银行重点行业清单",       "「人行八大高碳+我行主要行业」中的我行重点行业具体代码清单", "核算任务行业范围"],
        ["C13", "分行发起任务审批节点",       "分行发起任务时，分行审核通过是否还需要总行复核，还是直接进入排放计算", "数据审核流程"],
        ["C14", "数据采集归集合并是否分期",   "客户归集合并功能是否本期全量实施，还是分P0/P1/P2/P3分期实施（见改造计划文档）", "数据采集模块"],
    ])

    # ── 4.4 非功能性需求 ─────────────────────────────────────────────────────
    hp(doc, "4.4非功能性需求", 2)

    hp(doc, "4.4.1性能", 3)
    for item in [
        "系统应支持年度全行信贷台账量级的候选清单筛选、正式清单生成、数据采集跟踪和批量排放计算。",
        "数据层读写采用内存缓存，避免每次操作触发全量JSON序列化；数据迁移只执行一次。",
        "页面路由复用已挂载的SPA Shell，切换页面只刷新内容区域，不重建整页DOM。",
        "「调取接口数据」操作批量处理，单次读取完整数据后用Map索引进行O(1)查找。",
        "数据采集列表渲染预建supplement/candidate/calculation索引，避免每行多次遍历。",
        "【待确认】生产环境单任务最大台账笔数、并发用户数、报告导出最大记录数等性能指标待确认。",
    ]: lp(doc, item)

    hp(doc, "4.4.2安全性", 3)
    for item in [
        "三类角色（总行绿金部/分行负责人/客户经理）按宿主系统权限控制菜单和数据可见范围。",
        "总行绿金部：全行任务/审核记录/碳账户可见；可执行碳账户状态管理。",
        "分行负责人：仅辖内数据；可执行分行审核，不可变更碳账户状态。",
        "客户经理：仅访问客户经理任务清单和收集填报页；不可访问管理端功能。",
        "基础配置（参数字段库/方法模板/行业配置）默认不可见，通过权限管理配置控制。",
    ]: lp(doc, item)

    hp(doc, "4.4.3可靠性", 3)
    for item in [
        "接口失败支持重试；收集任务关键状态（暂存/提交/审核/退回/通过）可追溯。",
        "已通过收集数据被管理员退回时，须作废原审核记录并保留退回原因。",
        "格澜接口调取失败不阻断主流程，标记gelanStatus并允许重试或转其他路径。",
        "企业碳账户仅接收已确认结果的数据，保证账户内数据与任务确认结果一致。",
        "localStorage存储不超过4MB（含碳账户数据压缩）；配额满时自动压缩并提示。",
    ]: lp(doc, item)

    hp(doc, "4.4.4业务数据处理", 3)
    for item in [
        "月均贷款余额：系统从接口台账按12个月月末余额计算；部分月份有余额时按实有存续月份数计算均值，而非固定除以12（详见3.2节）。",
        "平均资产总额：上年末+当年末合并报表资产总额之和除以2（详见3.2.3节）；候选清单展示当年末资产，正式清单展示平均资产总额。",
        "行业识别优先级：候选清单筛选和业务种类识别均以贷款投向所属行业（gbIndustryCode）为主要依据；企业所属行业（industryMajor）作为辅助展示字段保留。",
        "客户归集合并：第3.4节为扩展设计方案；当前 HTML 原型按正式清单逐笔派发与跟踪，归因排放仍按笔计算。",
        "方法优先级：报告法 > 物理活动法-能源法 > 物理活动法-产品法 > 经济活动法 > 其他计算法；分行审核通过时须明确选定最终采用方法。",
        "数据采集与直算并行：即使格澜/经济法已获取排放底稿，仍须完整保留派发收集任务的业务动作。",
        "格澜接口数据规则：成功返回后写入gelanEntityEmission并同步到企业碳账户；预填至「其他来源报告法」Tab，来源文案固定为「格澜数据-各地区企业环境信息披露平台」，客户经理可编辑数值。",
    ]: lp(doc, item)

    hp(doc, "4.4.5设计约束", 3)
    for item in [
        "页面风格、菜单框架、权限控制、审批入口符合绿金系统现有规范。",
        "审核文案统一使用「退回」，不使用「驳回」；数据采集文案统一使用「收集」，不使用「补录」。",
        "基础配置模块（参数字段库/方法模板/行业配置）默认不在侧栏展示，通过权限管理页面开启。",
        "能源法、产品法等计算估算公式以业务确认版本为准，当前原型暂不自动试算。",
    ]: lp(doc, item)

    # ── 4.5 数据管理目标 ─────────────────────────────────────────────────────
    hp(doc, "4.5数据管理目标", 2)
    add_plain_table(doc, ["数据对象", "管理要求"], [
        ["核算任务",       "保留任务配置（年度/行业/组织范围/余额口径）、流程步骤、发起组织、确认结果时间。"],
        ["候选/正式清单", "保留台账字段（含月均余额、存续月份、平均资产）、纳入状态、业务种类、项目明细、锁定状态、归集单元关联。"],
        ["收集任务",       "保留归集单元成员列表、填报字段（含各方法数据）、附件元数据、提交状态、审核状态、退回原因、格澜预填数据。"],
        ["计算结果",       "保留方法、主体排放、归因排放、质量得分/等级、计算完成时间；格澜/直算/兜底来源标记；DQR综合值。"],
        ["报告记录",       "保留导出范围、模板、格式、记录数、总排放、生成时间和操作人。"],
        ["企业碳账户",     "账户档案、年度排放摘要（annualProfiles）、逐笔排放记录（carbonAccountRecords）、状态变更日志；编辑档案不覆盖核算结果。"],
        ["因子库",         "内置/自定义因子分离；来源说明留存；导入历史留痕；无版本年度概念。"],
        ["方法模板",       "发布状态（draft/published）变更记录；引用字段版本快照。"],
        ["行业配置",       "全量行业分类数据及标识变更记录；标识变更影响后续任务范围配置。"],
    ])
    pending(doc, "RTO/RPO指标、备份周期、日志保留周期按行内灾备规范确认。")

    np(doc, "")
    np(doc, "— 文档结束 —")
    np(doc, "版本 v0.62 · 2026.06.29 · 起草：Cursor · 待评审")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_doc()
