#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""根据截图生成用户操作手册 Word 文档"""

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SHOT_DIR = ROOT / 'docs' / 'manual' / 'screenshots'
MANUAL_DIR = ROOT / 'docs' / 'manual'
MANUAL_VERSION = 'v0.1'
MANUAL_DATE = '20260526'  # 定稿日期 YYYYMMDD
OUT_DOCX = MANUAL_DIR / f'华夏银行投融资碳核算-用户操作手册-{MANUAL_DATE}-{MANUAL_VERSION}.docx'


def set_doc_font(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_title(doc, text, level=0):
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')


def add_image_path(doc, path, caption, width_cm=15.5):
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
    else:
        doc.add_paragraph(f'[缺少图片: {path.name}]')


def add_image(doc, name, caption):
    add_image_path(doc, SHOT_DIR / f'{name}.png', caption)


def module_section(doc, title, intro, steps_with_shots):
    doc.add_heading(title, level=2)
    add_para(doc, intro)
    for idx, (step_text, shot_name, cap) in enumerate(steps_with_shots, 1):
        doc.add_paragraph(f'{idx}. {step_text}')
        if shot_name:
            add_image(doc, shot_name, cap)


def build():
    doc = Document()
    set_doc_font(doc)

    doc.core_properties.title = '华夏银行投融资碳核算-用户操作手册'
    doc.core_properties.subject = f'用户操作手册 {MANUAL_VERSION}'
    doc.core_properties.comments = f'版本 {MANUAL_VERSION}，定稿日期 {MANUAL_DATE}'

    add_title(doc, '华夏银行投融资碳核算系统')
    add_title(doc, '用户操作手册（总行 / 分行 / 客户经理）', level=1)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'文档版本：{MANUAL_VERSION}　　定稿日期：{MANUAL_DATE[:4]}-{MANUAL_DATE[4:6]}-{MANUAL_DATE[6:8]}')
    run.font.size = Pt(11)
    run.font.color.rgb = None
    add_para(doc, '编制说明：本手册基于绿金系统「投融资碳核算」演示原型编写，按钮名称以【】标注。')

    # 一、业务流程说明（文字描述）
    doc.add_heading('一、业务流程说明', level=1)
    add_para(doc, '投融资碳核算嵌入绿金系统，采用「总行定规则、分行做组织、客户经理补信息、系统做计算」的协同模式。审批流转在绿金系统待办中处理，碳核算模块提供任务、清单、补录与计算能力。')

    doc.add_heading('1.1 核算任务六步主流程', level=2)
    steps_main = [
        '范畴确定：创建核算任务，明确年度、行业范围、组织范围与输出目标。',
        '清单识别：同步信贷台账，按规则筛选候选业务，生成正式清单。',
        '对象边界：确认并锁定正式清单，明确核算对象类型与排放边界。',
        '数据采集：发放补录任务或执行经济法直算，组织客户经理填报并完成审核。',
        '排放计算：汇总主体排放与归因排放，确认计算结果。',
        '生成报告：按监管或管理口径导出 Excel/Word 报告。',
    ]
    for i, s in enumerate(steps_main, 1):
        doc.add_paragraph(f'（{i}）{s}')

    doc.add_heading('1.2 任务发起（两个起点）', level=2)
    add_para(doc, '核算任务可由总行绿金部或分行负责人发起，二选一：')
    add_para(doc, '（1）总行发起：在新建任务时选择【总行发起】，由总行组织全行或指定范围的清单识别、数据采集与后续计算。')
    add_para(doc, '（2）分行发起：选择【分行发起】并指定发起分行，由该分行组织辖内数据采集；清单识别及后续环节与总行发起共用同一套系统流程。')
    add_para(doc, '无论哪种发起方式，确认正式清单并锁定后，均进入数据采集、排放计算与报告生成等共用环节。')

    doc.add_heading('1.3 数据采集与补录协同', level=2)
    add_para(doc, '正式清单锁定后，发起方在「数据采集」页对必收数业务【发放补录任务】；对经济法直算类业务可【经济法直算】。客户经理在「数据补录」中【去填报】→【提交数据】→【提交审核】；分行通过「数据补录」看板统筹辖内进度。')

    doc.add_heading('1.4 数据审核（因发起方不同而不同）', level=2)
    add_para(doc, '（1）总行发起的任务：客户经理提交审核后，先由分行负责人【分行初审】；初审通过后，由总行绿金部【总行终审】；终审通过后进入排放计算。')
    add_para(doc, '（2）分行发起的任务：客户经理提交审核后，由分行负责人【分行终审】即可，无需总行终审；通过后直接进入排放计算。')

    doc.add_heading('1.5 各角色职责摘要', level=2)
    add_para(doc, '总行绿金部：可发起任务；清单确认与锁定；派发补录、经济法直算；总行发起任务时执行总行终审；排放计算确认与报告导出；排放因子库与接口管理。')
    add_para(doc, '分行负责人：可发起本分行业务；辖内补录组织与进度跟踪；总行发起任务时做分行初审，分行发起任务时做分行终审。')
    add_para(doc, '客户经理：接收本人名下补录任务；在线填报（五类核算方法择一）；提交数据与提交审核；查看审批进度。')

    doc.add_page_break()

    # 二、角色权限
    doc.add_heading('二、角色与菜单权限', level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    hdr = ['能力', '总行绿金部', '分行负责人', '客户经理']
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ('核算任务管理', '✓ 全行', '✓', '✗'),
        ('数据补录 / 在线填报', '✗', '✓ 本分行', '✓ 本人'),
        ('数据审核', '✓ 总行终审', '✓ 分行初审', '仅查看本人'),
        ('排放因子库', '✓', '✓', '✗'),
        ('接口管理', '✓', '✓', '✗'),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    doc.add_page_break()

    # 三、总行
    doc.add_heading('三、总行绿金部操作手册', level=1)

    sections_hq = [
        (
            '3.1 核算任务管理',
            '核算年度任务的创建与全生命周期管理入口，配置行业范围、组织范围、余额口径及输出目标。',
            [
                ('登录绿金系统，进入投融资碳核算，确认顶栏角色为「总行绿金部」，选择当前核算任务。', 'HQ-01-任务列表', '图 3-1 核算任务列表'),
                ('在筛选区输入条件后点击【查询】；需清空时点击【重置】。', 'HQ-02-任务筛选', '图 3-2 任务筛选'),
                ('点击【+ 新建核算任务】，填写必填项后点击【保存并启动】。', 'HQ-03-新建任务', '图 3-3 新建核算任务'),
                ('对已有任务可【编辑】、【查看】或【删除】（确认弹窗中点击【确认删除】）。', 'HQ-04-任务操作', '图 3-4 任务操作'),
            ],
        ),
        (
            '3.2 清单识别（候选业务清单）',
            '从信贷台账按指引规则筛选候选业务，支持调整筛选条件后生成正式清单。',
            [
                ('在任务流程条进入「清单识别」。', 'HQ-05-流程条-清单识别', '图 3-5 清单识别'),
                ('点击【从接口同步台账（×年度）】刷新候选数据。', 'HQ-06-同步台账', '图 3-6 同步台账'),
                ('勾选业务品种、主体类型、行业等条件，点击【查询】或【恢复默认筛选条件】。', 'HQ-07-候选筛选', '图 3-7 候选筛选'),
                ('核对后点击【生成正式清单（N 笔）】。', 'HQ-08-生成正式清单', '图 3-8 生成正式清单'),
            ],
        ),
        (
            '3.3 正式清单确认与对象边界',
            '将候选转为正式待核算清单并锁定，随后确认核算对象与边界。',
            [
                ('在「正式清单确认」页勾选记录，点击【确认锁定】。', 'HQ-09-正式清单', '图 3-9 正式清单'),
                ('锁定后状态为已确认。', 'HQ-10-确认锁定', '图 3-10 确认锁定'),
                ('打开「核算对象与边界」核对项目/主体、边界与控制法。', 'HQ-11-对象边界', '图 3-11 对象边界'),
            ],
        ),
        (
            '3.4 数据采集',
            '对正式清单发放补录任务或执行经济法直算，跟踪填报与审核进度。',
            [
                ('进入「数据采集」查看统计卡片。', 'HQ-12-数据采集总览', '图 3-12 数据采集总览'),
                ('勾选未派发记录，点击【发放补录任务】。', 'HQ-13-发放补录', '图 3-13 发放补录'),
                ('对经济法直算类记录点击【经济法直算】。', 'HQ-14-经济法直算', '图 3-14 经济法直算'),
                ('使用筛选【查询】/【重置】；可【查看填报】或【驳回】。', 'HQ-15-数据采集筛选与操作', '图 3-15 数据采集操作'),
                ('全部具备主体排放后点击【一键提交数据】。', 'HQ-16-一键提交数据', '图 3-16 一键提交数据'),
            ],
        ),
        (
            '3.5 数据审核（总行终审）',
            '总行发起任务时，补录经分行初审后进入总行终审（审批在绿金待办完成，本模块为演示入口）。',
            [
                ('侧栏进入【数据审核】。', 'HQ-17-审核列表', '图 3-17 审核列表'),
                ('待办点击【审核】，核对后【审核通过】或【审核不通过】并在弹窗【确认】。', 'HQ-18-总行审核', '图 3-18 总行审核'),
                ('已处理记录可【查看】。', 'HQ-19-审核查看', '图 3-19 审核查看'),
            ],
        ),
        (
            '3.6 碳排放计算与核算结果',
            '汇总主体排放、归因排放与分行业占比，确认后方可导出报告。',
            [
                ('进入「排放计算」查看汇总。', 'HQ-20-排放计算', '图 3-20 排放计算'),
                ('核对后点击【确认结果】。', 'HQ-21-确认结果', '图 3-21 确认结果'),
                ('打开核算结果页查看明细与 DQR。', 'HQ-22-核算结果', '图 3-22 核算结果'),
            ],
        ),
        (
            '3.7 生成报告',
            '按监管或管理口径导出 Excel/Word，并保留历史记录。',
            [
                ('进入「生成报告」，选择导出范围与模板。', 'HQ-23-报告配置', '图 3-23 报告配置'),
                ('点击【导出 Excel】或【导出 Word】，在历史记录【下载】。', 'HQ-24-报告导出与下载', '图 3-24 报告导出'),
                ('完成后可【返回核算任务列表】。', 'HQ-25-返回任务列表', '图 3-25 返回任务列表'),
            ],
        ),
        (
            '3.8 排放因子库',
            '维护五类核算方法对应的行业排放因子。',
            [
                ('侧栏【排放因子库】浏览列表。', 'HQ-26-因子列表', '图 3-26 因子列表'),
                ('【新增因子】或编辑页填写后【保存】。', 'HQ-27-因子维护', '图 3-27 因子维护'),
            ],
        ),
        (
            '3.9 接口管理',
            '监控信贷台账月度批次；失败可【重新获取】。',
            [
                ('顶栏点击【接口管理】。', 'HQ-28-接口列表', '图 3-28 接口列表'),
                ('失败批次可【重新获取】，成功批次【查看】。', 'HQ-29-接口重试', '图 3-29 接口管理'),
            ],
        ),
    ]

    for title, intro, steps in sections_hq:
        module_section(doc, title, intro, steps)

    doc.add_page_break()

    # 四、分行
    doc.add_heading('四、分行绿金负责人操作手册', level=1)

    sections_br = [
        (
            '4.1 数据补录（分行看板）',
            '辖内补录任务总览，按客户、客户经理、截止时间跟踪进度。',
            [
                ('顶栏切换「分行负责人」，侧栏【数据补录】。', 'BR-01-分行补录看板', '图 4-1 分行补录看板'),
                ('查看待处理/填报中/已完成/已退回统计。', 'BR-02-分行统计卡片', '图 4-2 统计卡片'),
                ('督促客户经理完成【去填报】等待办。', 'BR-03-分行列表操作', '图 4-3 列表操作'),
            ],
        ),
        (
            '4.2 数据审核（分行初审）',
            '审核辖内客户经理已提交的补录数据。',
            [
                ('侧栏【数据审核】。', 'BR-04-分行审核列表', '图 4-4 分行审核列表'),
                ('点击【审核】，【审核通过】或【审核不通过】并【确认】。', 'BR-05-分行初审', '图 4-5 分行初审'),
                ('【查看】跟踪已进入总行终审的记录。', 'BR-06-审核进度', '图 4-6 审核进度'),
            ],
        ),
        (
            '4.3 任务进度查看',
            '在核算任务中查看本分行相关进度。',
            [
                ('【核算任务管理】打开任务【查看】。', 'BR-07-分行查看任务', '图 4-7 查看任务'),
                ('进入「数据采集」筛选本分行客户。', 'BR-08-分行数据采集', '图 4-8 分行数据采集'),
            ],
        ),
    ]

    for title, intro, steps in sections_br:
        module_section(doc, title, intro, steps)

    doc.add_page_break()

    # 五、客户经理
    doc.add_heading('五、客户经理操作手册', level=1)
    add_para(doc, '客户经理侧栏仅含【数据补录】（任务清单与在线填报）。')

    sections_mgr = [
        (
            '5.1 客户经理任务清单',
            '展示派发给本人的补录待办及缺口字段数。',
            [
                ('顶栏切换「客户经理」，进入任务清单。', 'MGR-01-客户经理任务', '图 5-1 任务清单'),
                ('点击【去填报】处理待办，【查看进度】查看审核中记录。', 'MGR-02-任务列表操作', '图 5-2 列表操作'),
            ],
        ),
        (
            '5.2 在线补录填报',
            '按指引优先级在五种方法中择一填报并上传附件。',
            [
                ('【去填报】进入碳排放信息采集。', 'MGR-03-填报页概览', '图 5-3 填报概览'),
                ('补齐「企业基本信息」。', 'MGR-04-基本信息', '图 5-4 基本信息'),
                ('切换方法 Tab 填写排放数据并上传附件。', 'MGR-05-方法填报', '图 5-5 方法填报'),
                ('【暂存】保存草稿；【提交数据】完成填报。', 'MGR-06-暂存与提交数据', '图 5-6 提交数据'),
                ('返回列表后【提交审核】，弹窗【确认提交】。', 'MGR-07-提交审核与流程', '图 5-7 提交审核'),
                ('若被驳回，按提示【重新填报】。', 'MGR-08-驳回重填', '图 5-8 驳回重填'),
                ('完成后【返回】任务清单。', 'MGR-09-返回列表', '图 5-9 返回列表'),
            ],
        ),
    ]

    for title, intro, steps in sections_mgr:
        module_section(doc, title, intro, steps)

    doc.add_page_break()

    doc.add_heading('六、端到端演示路径（约 10 分钟）', level=1)
    add_para(doc, '1. 总行：新建/选择任务 → 同步台账 → 生成正式清单 → 确认锁定')
    add_para(doc, '2. 总行：数据采集 →【发放补录任务】/【经济法直算】')
    add_para(doc, '3. 客户经理：【去填报】→【提交数据】→【提交审核】')
    add_para(doc, '4. 分行：【数据审核】→【审核通过】')
    add_para(doc, '5. 总行：终审 →【确认结果】→ 报告【导出 Excel】')
    add_para(doc, '演示前可点击顶栏【重置数据】加载样例任务 T2025001。')

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print('已生成:', OUT_DOCX)


if __name__ == '__main__':
    build()
